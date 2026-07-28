import io
import zipfile

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

_VERDE = RGBColor(0x06, 0x5F, 0x46)

_ELEMENTOS = [
    ("oxigenacion",  1, "Apoyo para la oxigenación del paciente"),
    ("alimentacion", 2, "Apoyo para la alimentación del paciente"),
    ("eliminacion",  3, "Apoyo para la eliminación del paciente"),
    ("confort",      4, "Apoyo para el confort, higiene y descanso"),
    ("seguridad",    5, "Medidas de seguridad y protección del paciente"),
    ("postmortem",   6, "Cuidados post mortem"),
    ("autocuidado",  7, "Orientación para el autocuidado del paciente"),
]


def _titulo(doc: Document, texto: str) -> None:
    h = doc.add_heading("", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run(texto)
    r.font.color.rgb = _VERDE
    r.font.size = Pt(14)


def _encabezado_candidato(doc: Document, datos: dict) -> None:
    doc.add_paragraph(f"Candidato: {datos.get('cand16Nombre', '')} {datos.get('cand16Apellidos', '')}")
    doc.add_paragraph(f"CURP: {datos.get('cand16CURP', '')}")
    doc.add_paragraph(f"Unidad médica: {datos.get('cand16UnidadMedica', '')}")
    doc.add_paragraph(f"Centro de Evaluación: {datos.get('cand16CE', '')}")
    doc.add_paragraph(f"Evaluador: {datos.get('cand16Evaluador', '')}")
    doc.add_paragraph(f"Fecha de evaluación: {datos.get('cand16Fecha', '')}")
    doc.add_paragraph()


def _docx_bytes(doc: Document) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def generar_portada(payload: dict) -> bytes:
    doc = Document()
    _titulo(doc, "Portafolio de Evidencias")
    _titulo(doc, "Norma EC0616 — Auxiliar de Enfermería")
    doc.add_paragraph()
    datos = payload.get("datos", {})
    _encabezado_candidato(doc, datos)
    doc.add_paragraph()
    doc.add_heading("Índice de Evidencias", level=2)
    for elem_id, num, titulo in _ELEMENTOS:
        doc.add_paragraph(f"Elemento {num}: {titulo}", style="List Bullet")
    doc.add_paragraph("Declaratoria del candidato", style="List Bullet")
    return _docx_bytes(doc)


def _generar_evidencia_elemento(payload: dict, elem_id: str, num: int, titulo: str) -> bytes:
    doc = Document()
    _titulo(doc, f"Evidencia — Elemento {num}")
    _titulo(doc, titulo)
    doc.add_paragraph("Norma EC0616 — Auxiliar de Enfermería")
    doc.add_paragraph()
    _encabezado_candidato(doc, payload.get("datos", {}))
    elem = payload.get(elem_id, {})
    doc.add_heading("Descripción de desempeños", level=2)
    for linea in (elem.get("desempenos", "") or "").split("\n"):
        doc.add_paragraph(linea)
    if elem.get("productos"):
        doc.add_heading("Productos / Evidencias documentales", level=2)
        for linea in elem.get("productos", "").split("\n"):
            doc.add_paragraph(linea)
    doc.add_paragraph()
    doc.add_paragraph("Firma del candidato: ___________________")
    doc.add_paragraph("Firma del evaluador: ___________________")
    return _docx_bytes(doc)


def generar_declaratoria(payload: dict) -> bytes:
    doc = Document()
    _titulo(doc, "Declaratoria del Candidato")
    _titulo(doc, "Norma EC0616 — Auxiliar de Enfermería")
    doc.add_paragraph()
    datos = payload.get("datos", {})
    nombre = f"{datos.get('cand16Nombre', '')} {datos.get('cand16Apellidos', '')}".strip()
    doc.add_paragraph(
        f"Yo, {nombre}, con CURP {datos.get('cand16CURP', '')}, declaro que la información "
        f"contenida en el presente portafolio de evidencias es verídica y corresponde a mi "
        f"desempeño real como auxiliar de enfermería en {datos.get('cand16UnidadMedica', '')}, "
        f"de conformidad con los criterios del estándar de competencia EC0616 del CONOCER."
    )
    doc.add_paragraph()
    doc.add_paragraph(f"Lugar y fecha: _________________, {datos.get('cand16Fecha', '')}")
    doc.add_paragraph()
    doc.add_paragraph("Firma del candidato: ___________________")
    doc.add_paragraph(f"Nombre: {nombre}")
    doc.add_paragraph(f"CURP: {datos.get('cand16CURP', '')}")
    doc.add_paragraph()
    doc.add_paragraph(f"Firma del evaluador: ___________________")
    doc.add_paragraph(f"Evaluador: {datos.get('cand16Evaluador', '')}")
    doc.add_paragraph(f"Centro de Evaluación: {datos.get('cand16CE', '')}")
    return _docx_bytes(doc)


def generar_zip_ec0616(payload: dict) -> bytes:
    docs = [("00_Portada_Portafolio.docx", generar_portada(payload))]
    for elem_id, num, titulo in _ELEMENTOS:
        filename = f"0{num}_Evidencia_{titulo.split()[0].replace(',','')}.docx"
        docs.append((filename, _generar_evidencia_elemento(payload, elem_id, num, titulo)))
    docs.append(("08_Declaratoria_Candidato.docx", generar_declaratoria(payload)))

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
        for nombre, contenido in docs:
            zf.writestr(nombre, contenido)
    buf.seek(0)
    return buf.read()
