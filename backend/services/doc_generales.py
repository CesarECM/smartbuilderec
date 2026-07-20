import math

from docx import Document

from services.doc_helpers import tabla_encabezado, _docx_bytes, _titulo


def generar_contrato_aprendizaje(data) -> bytes:
    doc = Document()
    _titulo(doc, "CONTRATO DE APRENDIZAJE")
    tabla_encabezado(doc, data.datos)

    doc.add_heading("Alcance e instrucciones del contrato de aprendizaje", level=2)
    doc.add_paragraph(
        "El presente contrato de aprendizaje establece los compromisos mutuos entre "
        "el instructor y los participantes para garantizar el logro de los objetivos "
        "del curso. Ambas partes se comprometen a cumplir con lo acordado durante el "
        "desarrollo de la sesión."
    )

    doc.add_paragraph()

    doc.add_heading("Compromisos del Instructor", level=2)
    compromisos_instructor = [
        "Me comprometo a que el curso sea dinámico y práctico.",
        "Me comprometo que los contenidos sean comprensibles y de fácil asimilación.",
        "Me comprometo a que se cumplan los objetivos del curso.",
        "Me comprometo a poder cumplir las expectativas del curso.",
        "Me comprometo a aclarar dudas respetando los tiempos del curso.",
        "Me comprometo a dejarte algo práctico para su vida profesional y laboral.",
    ]
    for i, c in enumerate(compromisos_instructor, 1):
        doc.add_paragraph(f"{i}. {c}", style="List Number")

    doc.add_paragraph()

    doc.add_heading("Compromisos del Participante", level=2)
    for i in range(1, 7):
        doc.add_paragraph(f"{i}. _______________________________________________", style="List Number")

    doc.add_paragraph()
    doc.add_paragraph("_________________________________          _______________________________")
    doc.add_paragraph("Nombre y Firma del Participante                        Firma del Instructor")

    return _docx_bytes(doc)


def generar_lista_requerimientos(data) -> bytes:
    doc = Document()
    _titulo(doc, "LISTA DE VERIFICACIÓN DE REQUERIMIENTOS")
    tabla_encabezado(doc, data.datos, sin_participante=True)

    mat = data.materiales or {}
    participantes = data.datos.participantes or 0

    secciones = [
        ("Instalaciones, mobiliario y su distribución",        mat.get("instalaciones", "")),
        ("Equipo de apoyo",                                    mat.get("equipo", "")),
        ("Materiales didácticos de apoyo y servicios",         mat.get("materialesDidacticos", "")),
        ("Requerimientos humanos",                             mat.get("humanos", "")),
        ("Otros requerimientos",                               mat.get("otros", "")),
    ]

    _ITEMS_SEGURIDAD = [
        "Salidas de emergencia señalizadas y despejadas",
        "Extintor accesible, vigente y visible",
        "Botiquín de primeros auxilios disponible",
        "Medidas de higiene (desinfectante, ventilación, distancia entre participantes)",
        "Protocolo de evacuación conocido por todos los participantes",
    ]

    def _items(texto: str) -> list:
        if not texto:
            return []
        items = []
        for linea in texto.split("\n"):
            limpia = linea.strip().lstrip("•-–*").strip()
            if limpia:
                items.append(limpia)
        return items

    for titulo, contenido in secciones:
        doc.add_heading(titulo, level=2)
        items = _items(contenido)
        num_filas_datos = max(len(items), 3)
        t = doc.add_table(rows=num_filas_datos + 1, cols=5)
        t.style = "Table Grid"
        encabezados = ["No.", "Descripción", "Cantidad", "Existe", "No existe"]
        for j, enc in enumerate(encabezados):
            t.rows[0].cells[j].paragraphs[0].add_run(enc).bold = True
        for i, item in enumerate(items, start=1):
            t.rows[i].cells[0].text = str(i)
            t.rows[i].cells[1].text = item
        for i in range(len(items) + 1, num_filas_datos + 1):
            t.rows[i].cells[0].text = str(i)
        doc.add_paragraph()

    # Sección de seguridad: ítems mínimos hardcoded + ítems del usuario
    doc.add_heading("Medidas de Salud / Seguridad / Higiene / Protección Civil", level=2)
    items_seg = _ITEMS_SEGURIDAD + [i for i in _items(mat.get("seguridad", "")) if i not in _ITEMS_SEGURIDAD]
    t = doc.add_table(rows=len(items_seg) + 1, cols=3)
    t.style = "Table Grid"
    for j, enc in enumerate(["No.", "Ítem de verificación", "Verificado ✓"]):
        t.rows[0].cells[j].paragraphs[0].add_run(enc).bold = True
    for i, item in enumerate(items_seg, start=1):
        t.rows[i].cells[0].text = str(i)
        t.rows[i].cells[1].text = item
    doc.add_paragraph()

    # Formatos individuales por participante
    doc.add_heading("Formatos individuales (uno por participante)", level=2)
    hojas_asistencia = math.ceil(participantes / 20) if participantes > 0 else 1
    formatos_ind = [
        ("Evaluación Diagnóstica",          participantes),
        ("Contrato de Aprendizaje",         participantes),
        ("Evaluación Formativa",            participantes),
        ("Evaluación Sumativa",             participantes),
        ("Evaluación de Reacción",          participantes),
        ("Lista de Asistencia (hojas)",     hojas_asistencia),
    ]
    t = doc.add_table(rows=len(formatos_ind) + 1, cols=5)
    t.style = "Table Grid"
    for j, enc in enumerate(["No.", "Formato", "Cantidad", "Listo", "Observaciones"]):
        t.rows[0].cells[j].paragraphs[0].add_run(enc).bold = True
    for i, (nombre, cantidad) in enumerate(formatos_ind, start=1):
        t.rows[i].cells[0].text = str(i)
        t.rows[i].cells[1].text = nombre
        t.rows[i].cells[2].text = str(cantidad) if participantes > 0 else "—"
    doc.add_paragraph()

    return _docx_bytes(doc)
