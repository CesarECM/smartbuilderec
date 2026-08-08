# ─── services/doc_ec0616_tutor_helpers.py — Generación DOCX Modo Tutor EC0616 ──

import io
import zipfile

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from services.pdf_ec0616_tutor_helpers import generar_pdf_ec0616_tutor

_VERDE = RGBColor(0x06, 0x5F, 0x46)

_ELEM_LABELS = [
    (1, "Apoyo para la oxigenación del paciente"),
    (2, "Apoyo para la alimentación del paciente"),
    (3, "Apoyo para la eliminación del paciente"),
    (4, "Apoyo para el confort, higiene y descanso"),
    (5, "Medidas de seguridad y protección del paciente"),
    (6, "Cuidados post mortem"),
    (7, "Orientación para el autocuidado del paciente"),
]

_ENCUESTA_LABELS = {
    "atencion_evaluador":     "Atención del evaluador",
    "claridad_instrucciones": "Claridad de las instrucciones",
    "instalaciones":          "Condiciones de las instalaciones",
    "tiempo_adecuado":        "Tiempo de evaluación adecuado",
}

# ── Helpers ───────────────────────────────────────────────────────────────────

def _titulo(doc: Document, texto: str) -> None:
    h = doc.add_heading("", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run(texto)
    r.font.color.rgb = _VERDE
    r.font.size = Pt(13)


def _seccion(doc: Document, texto: str) -> None:
    h = doc.add_heading("", level=2)
    h.add_run(texto).font.size = Pt(11)


def _fila(doc: Document, label: str, value) -> None:
    doc.add_paragraph(f"{label}: {value or '—'}")


def _docx_bytes(doc: Document) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _encabezado(doc: Document, ficha: dict) -> None:
    nombre = f"{ficha.get('nombre_candidato', '')} {ficha.get('apellidos_candidato', '')}".strip()
    _fila(doc, "Candidato", nombre or "—")
    _fila(doc, "CURP", ficha.get("curp"))
    _fila(doc, "Centro de Evaluación", ficha.get("nombre_ce"))
    _fila(doc, "Evaluador", ficha.get("nombre_evaluador"))
    _fila(doc, "Fecha de evaluación", ficha.get("fecha_evaluacion"))
    doc.add_paragraph()


# ── Doc 1: Ficha de Registro ──────────────────────────────────────────────────

def generar_ficha(ficha: dict, conf: dict) -> bytes:
    doc = Document()
    _titulo(doc, "Ficha de Registro del Candidato")
    _titulo(doc, "Norma EC0616 — Auxiliar de Enfermería")
    doc.add_paragraph()

    _seccion(doc, "Datos Personales")
    nombre = f"{ficha.get('nombre_candidato', '')} {ficha.get('apellidos_candidato', '')}".strip()
    _fila(doc, "Nombre completo", nombre)
    _fila(doc, "CURP", ficha.get("curp"))
    _fila(doc, "Fecha de nacimiento", ficha.get("fecha_nacimiento"))
    _fila(doc, "Género", ficha.get("genero"))
    _fila(doc, "Lugar de nacimiento", ficha.get("lugar_nacimiento"))
    _fila(doc, "Nacionalidad", ficha.get("nacionalidad"))

    _seccion(doc, "Domicilio")
    dom = " ".join(filter(None, [
        ficha.get("domicilio_calle"), ficha.get("domicilio_numero"),
        ficha.get("domicilio_colonia"), ficha.get("domicilio_cp"),
        ficha.get("domicilio_ciudad"), ficha.get("domicilio_entidad"),
    ]))
    _fila(doc, "Domicilio", dom or None)
    _fila(doc, "Correo electrónico", ficha.get("email"))
    _fila(doc, "Teléfono", ficha.get("telefono"))

    _seccion(doc, "Datos de la Evaluación")
    _fila(doc, "Centro de Evaluación", ficha.get("nombre_ce"))
    _fila(doc, "Evaluador", ficha.get("nombre_evaluador"))
    _fila(doc, "Fecha de evaluación", ficha.get("fecha_evaluacion"))
    _fila(doc, "Consentimiento RENAP", ficha.get("consentimiento_renap"))

    _seccion(doc, "Información Complementaria")
    _fila(doc, "Sabe leer y escribir", conf.get("leer_escribir"))
    _fila(doc, "Cuenta con estudios", conf.get("tiene_estudios"))
    if conf.get("estudios_cuales"):
        _fila(doc, "  Tipo de estudios", conf.get("estudios_cuales"))
        _fila(doc, "  Último grado", conf.get("ultimo_grado"))
    _fila(doc, "Trabaja actualmente", conf.get("trabaja_actualmente"))
    if conf.get("puesto_trabajo"):
        _fila(doc, "  Puesto", conf.get("puesto_trabajo"))
        _fila(doc, "  Empresa/institución", conf.get("empresa_trabajo"))
    _fila(doc, "Discapacidad", conf.get("tiene_discapacidad"))
    if conf.get("discapacidad_tipo"):
        _fila(doc, "  Tipo", ", ".join(conf.get("discapacidad_tipo") or []))
    _fila(doc, "Idiomas", conf.get("idiomas"))
    _fila(doc, "Certificación previa", conf.get("tiene_certificacion"))
    if conf.get("certificacion_cuales"):
        _fila(doc, "  Cuáles", conf.get("certificacion_cuales"))

    doc.add_paragraph()
    doc.add_paragraph("Firma del candidato: ___________________")
    doc.add_paragraph("Firma del evaluador: ___________________")
    return _docx_bytes(doc)


# ── Doc 2: Diagnóstico ────────────────────────────────────────────────────────

def generar_diagnostico(ficha: dict, diag: dict) -> bytes:
    doc = Document()
    _titulo(doc, "Resultado del Diagnóstico")
    _titulo(doc, "Norma EC0616 — Auxiliar de Enfermería")
    doc.add_paragraph()
    _encabezado(doc, ficha)
    _fila(doc, "Reactivos correctos", f"{diag.get('correctas', 0)} de {diag.get('total', 0)}")
    _fila(doc, "Posibilidad de éxito", diag.get("posibilidad"))
    _fila(doc, "Sugerencia", diag.get("sugerencia"))
    _fila(doc, "Fecha del diagnóstico", diag.get("fecha"))
    return _docx_bytes(doc)


# ── Doc 3: Plan de Evaluación ─────────────────────────────────────────────────

def generar_plan(ficha: dict, plan: dict, plan_config: dict) -> bytes:
    doc = Document()
    _titulo(doc, "Plan de Evaluación")
    _titulo(doc, "Norma EC0616 — Auxiliar de Enfermería")
    doc.add_paragraph()
    _encabezado(doc, ficha)

    _seccion(doc, "Logística")
    _fila(doc, "Fecha de evaluación", plan.get("fecha_eval") or ficha.get("fecha_evaluacion"))
    _fila(doc, "Horario", plan.get("hora"))
    _fila(doc, "Lugar", plan.get("lugar"))
    _fila(doc, "Número de sesiones", plan.get("sesiones"))
    _fila(doc, "Fecha de resultados", plan.get("fecha_resultados"))
    _fila(doc, "Medio de comunicación", plan.get("lugar_resultados"))
    _fila(doc, "Observaciones", plan.get("observaciones"))

    actividades = (plan_config or {}).get("actividades") or []
    if actividades:
        _seccion(doc, f"Actividades a Evaluar ({len(actividades)})")
        for a in actividades:
            doc.add_paragraph(
                f"{a.get('num', '')}. [{a.get('tecnica', '')}] {a.get('descripcion', '')}",
                style="List Bullet",
            )
    return _docx_bytes(doc)


# ── Doc 4: IEC Aplicado ───────────────────────────────────────────────────────

def generar_iec_aplicado(ficha: dict, iec_sections: dict, iec_reactivos: list) -> bytes:
    doc = Document()
    _titulo(doc, "Instrumento de Evaluación de Competencias — Aplicado")
    _titulo(doc, "Norma EC0616 — Auxiliar de Enfermería")
    doc.add_paragraph()
    _encabezado(doc, ficha)

    marcas: dict = {}
    for i in range(1, 8):
        for r in (iec_sections.get(f"iec{i}", {}).get("reactivos") or []):
            marcas[r["codigo"]] = r.get("cumple")

    for num, largo in _ELEM_LABELS:
        _seccion(doc, f"Elemento {num}: {largo}")
        elem = [r for r in iec_reactivos if (r.get("codigo") or "").startswith(f"E{num}")]
        if not elem:
            doc.add_paragraph("Sin reactivos registrados.")
        for r in elem:
            cumple = marcas.get(r["codigo"])
            res = "C" if cumple is True else "NC" if cumple is False else "—"
            ahv = " [AHV]" if r.get("es_ahv") else ""
            p = doc.add_paragraph()
            run_res = p.add_run(f"[{res}] ")
            run_res.bold = True
            p.add_run(
                f"{r.get('codigo', '')}{ahv} — {r.get('descripcion', '')} "
                f"(Peso: {r.get('peso_relativo', '—')})"
            )
        doc.add_paragraph()
    return _docx_bytes(doc)


# ── Doc 5: Cédula de Evaluación ───────────────────────────────────────────────

def generar_cedula(ficha: dict, cierre: dict) -> bytes:
    doc = Document()
    _titulo(doc, "Cédula de Evaluación")
    _titulo(doc, "Norma EC0616 — Auxiliar de Enfermería")
    doc.add_paragraph()
    _encabezado(doc, ficha)

    juicio_raw = cierre.get("juicio", "—")
    juicio_label = (
        "Competente (C)" if juicio_raw == "C"
        else "Todavía No Competente (TNC)" if juicio_raw == "TNC"
        else juicio_raw
    )
    _fila(doc, "Puntaje IEC total", str(cierre.get("puntaje_total", "—")))
    _fila(doc, "Juicio de evaluación", juicio_label)
    _fila(doc, "Observaciones del evaluador", cierre.get("obs"))
    doc.add_paragraph()
    _fila(doc, "Nombre del evaluador", cierre.get("evaluador") or ficha.get("nombre_evaluador"))
    _fila(doc, "Fecha", cierre.get("fecha") or ficha.get("fecha_evaluacion"))
    doc.add_paragraph()
    doc.add_paragraph("Firma del evaluador: ___________________")
    doc.add_paragraph("Firma del candidato: ___________________")
    return _docx_bytes(doc)


# ── Doc 6: Encuesta de Satisfacción ──────────────────────────────────────────

def generar_encuesta(ficha: dict, cierre: dict) -> bytes:
    doc = Document()
    _titulo(doc, "Encuesta de Satisfacción")
    _titulo(doc, "Norma EC0616 — Auxiliar de Enfermería")
    doc.add_paragraph()
    _encabezado(doc, ficha)
    _seccion(doc, "Calificación (1 = deficiente  /  5 = excelente)")
    enc = cierre.get("encuesta") or {}
    for key, label in _ENCUESTA_LABELS.items():
        val = enc.get(key)
        _fila(doc, label, str(val) if val is not None else None)
    doc.add_paragraph()
    _fila(doc, "Comentarios libres", cierre.get("comentarios"))
    return _docx_bytes(doc)


# ── ZIP final ─────────────────────────────────────────────────────────────────

def generar_zip_ec0616_tutor(payload: dict, iec_reactivos: list, plan_config: dict) -> bytes:
    ficha  = payload.get("ficha", {})
    conf   = payload.get("confidencial", {})
    diag   = payload.get("diagnostico", {})
    plan   = payload.get("plan", {})
    cierre = payload.get("cierre", {})
    iec_sections = {f"iec{i}": payload.get(f"iec{i}", {}) for i in range(1, 8)}

    archivos = [
        ("01_Ficha_de_Registro.docx",    generar_ficha(ficha, conf)),
        ("02_Diagnostico.docx",           generar_diagnostico(ficha, diag)),
        ("03_Plan_de_Evaluacion.docx",    generar_plan(ficha, plan, plan_config)),
        ("04_IEC_Aplicado.docx",          generar_iec_aplicado(ficha, iec_sections, iec_reactivos)),
        ("05_Cedula_de_Evaluacion.docx",  generar_cedula(ficha, cierre)),
        ("06_Encuesta_Satisfaccion.docx", generar_encuesta(ficha, cierre)),
        ("portafolio_completo.pdf",       generar_pdf_ec0616_tutor(payload, iec_reactivos)),
    ]

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
        for nombre, contenido in archivos:
            zf.writestr(nombre, contenido)
    buf.seek(0)
    return buf.read()
