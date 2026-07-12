from docx.shared import Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

AZUL       = "1F3B6D"
AZUL_CLARO = "D6E4F0"
GRIS_CLARO = "F2F2F2"
AMARILLO   = "FFF9E6"
BLANCO     = "FFFFFF"
NEGRO      = "000000"
TW         = 10800  # ancho útil carta en twips

# ─── XML helpers ──────────────────────────────────────────────────────────────

def _shading(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    for e in tcPr.findall(qn('w:shd')): tcPr.remove(e)
    tcPr.append(shd)

def _borders(cell, color=AZUL, sz=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(sz))
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tcB.append(b)
    for e in tcPr.findall(qn('w:tcBorders')): tcPr.remove(e)
    tcPr.append(tcB)

def _pad(cell, top=80, bottom=80, left=120, right=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcM = OxmlElement('w:tcMar')
    for side, val in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), str(val))
        m.set(qn('w:type'), 'dxa')
        tcM.append(m)
    for e in tcPr.findall(qn('w:tcMar')): tcPr.remove(e)
    tcPr.append(tcM)

def _cw(cell, width):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    w = OxmlElement('w:tcW')
    w.set(qn('w:w'), str(width))
    w.set(qn('w:type'), 'dxa')
    for e in tcPr.findall(qn('w:tcW')): tcPr.remove(e)
    tcPr.append(w)

def _table_setup(table, col_widths):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(sum(col_widths)))
    tblW.set(qn('w:type'), 'dxa')
    for e in tblPr.findall(qn('w:tblW')): tblPr.remove(e)
    tblPr.append(tblW)
    lay = OxmlElement('w:tblLayout')
    lay.set(qn('w:type'), 'fixed')
    for e in tblPr.findall(qn('w:tblLayout')): tblPr.remove(e)
    tblPr.append(lay)
    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is None:
        tblGrid = OxmlElement('w:tblGrid')
        tbl.insert(1, tblGrid)
    else:
        for gc in tblGrid.findall(qn('w:gridCol')): tblGrid.remove(gc)
    for w in col_widths:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(w))
        tblGrid.append(gc)

# ─── Helpers de párrafo ───────────────────────────────────────────────────────

def _run(para, text, bold=False, color=NEGRO, size_pt=10):
    run = para.add_run(str(text or ''))
    run.font.name = 'Arial'
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    return run

def _clear(cell):
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)

def _para(cell, text='', bold=False, color=NEGRO, size_pt=10,
          align=None, sp_before=16, sp_after=16):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Twips(sp_before)
    p.paragraph_format.space_after  = Twips(sp_after)
    if align: p.alignment = align
    if text: _run(p, text, bold=bold, color=color, size_pt=size_pt)
    return p

def _lines(cell, text, size_pt=10, bold=False, color=NEGRO):
    _clear(cell)
    for line in str(text or '').split('\n'):
        _para(cell, line, bold=bold, color=color, size_pt=size_pt)

# ─── Tipos de celda ──────────────────────────────────────────────────────────

def _hdr(cell, text, width=None):
    _borders(cell); _pad(cell, 60, 60, 120, 120); _shading(cell, AZUL)
    if width: _cw(cell, width)
    _clear(cell)
    _para(cell, text, bold=True, color=BLANCO, size_pt=11,
          align=WD_ALIGN_PARAGRAPH.CENTER, sp_before=24, sp_after=24)

def _subhdr(cell, text, width=None):
    _borders(cell); _pad(cell); _shading(cell, AZUL_CLARO)
    if width: _cw(cell, width)
    _clear(cell)
    _para(cell, text, bold=True, color=AZUL,
          align=WD_ALIGN_PARAGRAPH.CENTER, sp_before=20, sp_after=20)

def _lbl(cell, text, width=None):
    _borders(cell); _pad(cell); _shading(cell, AZUL_CLARO)
    if width: _cw(cell, width)
    _clear(cell)
    _para(cell, text, bold=True, color=AZUL, sp_before=20, sp_after=20)

def _val(cell, text, width=None, size_pt=10, shade=None):
    _borders(cell); _pad(cell)
    if shade: _shading(cell, shade)
    if width: _cw(cell, width)
    _lines(cell, text, size_pt=size_pt)

def _val_s(cell, text, width=None):
    _val(cell, text, width=width, size_pt=9)

# ─── Utilidades de datos ──────────────────────────────────────────────────────

def _fmt_dur(minutos):
    m = int(minutos or 0)
    if m >= 60:
        h, r = divmod(m, 60)
        return f"{h} h {r} min" if r else f"{h} hrs"
    return f"{m} min"

def _tiempo(d, titulo):
    norm = lambda s: ' '.join(str(s or '').lower().split())
    for b in (d.get('tiempos') or []):
        for f in (b.get('filas') or []):
            if norm(f.get('titulo', '')) == norm(titulo):
                return f"{f.get('tiempo', '')} min"
    return ""

def _subtotal(d, seccion):
    for b in (d.get('tiempos') or []):
        if b.get('seccion') == seccion:
            return f"{sum(int(f.get('tiempo') or 0) for f in (b.get('filas') or []))} min"
    return "—"

def _inst(texto, fallback=""):
    if not texto: return fallback
    lines = [l.strip() for l in str(texto).split('\n') if l.strip()]
    return lines[0] if lines and len(lines[0]) <= 120 else fallback
