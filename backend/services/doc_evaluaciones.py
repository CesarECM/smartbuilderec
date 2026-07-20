from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from services.doc_helpers import tabla_encabezado, _docx_bytes, _titulo


def generar_evaluacion_diagnostica(data) -> bytes:
    doc = Document()

    titulo = doc.add_heading("EVALUACIÓN DIAGNÓSTICA", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitulo = doc.add_heading("CUESTIONARIO", level=2)
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    tabla_encabezado(doc, data.datos)

    doc.add_heading("Instrucciones de aplicación", level=2)
    instrucciones = [
        ("Alcance/beneficio:", "El presente cuestionario tiene como propósito conocer los conocimientos previos del participante antes de iniciar el curso."),
        ("Indicaciones para el facilitador:", "Entregue el cuestionario al participante antes de iniciar el curso. No hay respuestas correctas o incorrectas; el objetivo es identificar el nivel de conocimiento previo."),
        ("Condiciones de aplicación:", "Se aplicará al inicio del curso, de manera individual y sin consultas documentales."),
        ("Tiempo para desarrollar la actividad:", "Considerar el tiempo establecido en el documento de planeación."),
        ("Valor:", f"{data.evaluaciones.pctDiagnostica or data.evaluaciones.pctDiag or 0}% de la calificación total."),
    ]
    for etiqueta, texto in instrucciones:
        p = doc.add_paragraph()
        p.add_run(etiqueta).bold = True
        p.add_run(f" {texto}")

    doc.add_paragraph()

    reactivos_texto = (
        data.evaluaciones.instDiagnostica
        or data.evaluaciones.instDiag
        or ""
    ).strip()

    header_diag = (getattr(data.evaluaciones, "instDiagnosticaHeader", "") or "").strip()
    if header_diag:
        doc.add_paragraph(header_diag)

    if reactivos_texto:
        doc.add_heading("Cuestionario", level=2)
        for linea in reactivos_texto.split("\n"):
            if linea.strip():
                doc.add_paragraph(linea.strip())
    else:
        doc.add_heading("Cuestionario", level=2)
        for i in range(1, 11):
            doc.add_paragraph(f"{i}. __________________________________________________________")
            doc.add_paragraph()

    doc.add_paragraph()
    doc.add_paragraph("Nombre y firma del Participante: ________________________________________")

    return _docx_bytes(doc)


def generar_evaluacion_formativa_cotejo(data) -> bytes:
    doc = Document()
    _titulo(doc, "EVALUACIÓN FORMATIVA", "LISTA DE COTEJO")
    tabla_encabezado(doc, data.datos)

    doc.add_heading("Instrucciones de aplicación", level=2)
    instrucciones = [
        ("Alcance/beneficio:", "La presente lista de cotejo tiene como propósito identificar la comprensión y avance logrado por el participante."),
        ("Indicaciones para el facilitador:", "Revise el documento entregado por el participante y marque con una X en SÍ cuando cumple con los reactivos solicitados, y en NO en caso contrario."),
        ("Condiciones de aplicación:", "Se aplicará en el desarrollo del curso durante la técnica demostrativa."),
        ("Tiempo para desarrollar la actividad:", "Considerar el mismo tiempo establecido en el documento de planeación."),
        ("Valor:", f"{data.evaluaciones.pctFormativa or data.evaluaciones.pctForm or 0}% de la calificación total."),
    ]
    for etiqueta, texto in instrucciones:
        p = doc.add_paragraph()
        p.add_run(etiqueta).bold = True
        p.add_run(f" {texto}")

    header_form = (getattr(data.evaluaciones, "instFormativaHeader", "") or "").strip()
    if header_form:
        doc.add_paragraph(header_form)

    doc.add_paragraph()

    reactivos_texto = (
        data.evaluaciones.instFormativa
        or data.evaluaciones.instForm
        or ""
    ).strip()
    if reactivos_texto:
        doc.add_heading("Criterios de evaluación", level=2)
        for linea in reactivos_texto.split("\n"):
            if linea.strip():
                doc.add_paragraph(linea.strip())
        doc.add_paragraph()

    t = doc.add_table(rows=7, cols=5)
    t.style = "Table Grid"
    encabezados = ["Reactivo", "Descripción", "SÍ", "NO", "Observaciones"]
    for j, enc in enumerate(encabezados):
        t.rows[0].cells[j].paragraphs[0].add_run(enc).bold = True
    for i in range(1, 7):
        t.rows[i].cells[0].text = str(i)

    return _docx_bytes(doc)


def generar_evaluacion_formativa_guia(data) -> bytes:
    doc = Document()
    _titulo(doc, "EVALUACIÓN FORMATIVA", "GUÍA DE OBSERVACIÓN")
    tabla_encabezado(doc, data.datos)

    doc.add_heading("Instrucciones de aplicación", level=2)
    instrucciones = [
        ("Alcance/beneficio:", "La presente guía de observación tiene como propósito identificar la comprensión y correcto desempeño del participante."),
        ("Indicaciones para el facilitador:", "Observe cuidadosamente la ejecución de las actividades y marque con una X en SÍ cuando cumple con los reactivos solicitados, y en NO en caso contrario."),
        ("Condiciones de aplicación:", "Se aplicará en el desarrollo del curso durante la técnica demostrativa."),
        ("Tiempo para desarrollar la actividad:", "Considerar el mismo tiempo establecido en el documento de planeación."),
        ("Valor:", f"{data.evaluaciones.pctFormativa or data.evaluaciones.pctForm or 0}% de la calificación total."),
    ]
    for etiqueta, texto in instrucciones:
        p = doc.add_paragraph()
        p.add_run(etiqueta).bold = True
        p.add_run(f" {texto}")

    header_form = (getattr(data.evaluaciones, "instFormativaHeader", "") or "").strip()
    if header_form:
        doc.add_paragraph(header_form)

    doc.add_paragraph()

    reactivos_texto = (
        data.evaluaciones.instFormativa
        or data.evaluaciones.instForm
        or ""
    ).strip()
    if reactivos_texto:
        doc.add_heading("Criterios de evaluación", level=2)
        for linea in reactivos_texto.split("\n"):
            if linea.strip():
                doc.add_paragraph(linea.strip())
        doc.add_paragraph()

    t = doc.add_table(rows=7, cols=5)
    t.style = "Table Grid"
    encabezados = ["Reactivo", "Descripción", "SÍ", "NO", "Observaciones"]
    for j, enc in enumerate(encabezados):
        t.rows[0].cells[j].paragraphs[0].add_run(enc).bold = True
    for i in range(1, 7):
        t.rows[i].cells[0].text = str(i)

    return _docx_bytes(doc)


def generar_evaluacion_sumativa(data) -> bytes:
    doc = Document()
    _titulo(doc, "EVALUACIÓN SUMATIVA", "CUESTIONARIO")
    tabla_encabezado(doc, data.datos)

    doc.add_heading("Instrucciones de aplicación", level=2)
    instrucciones = [
        ("Alcance/beneficio:", "El presente cuestionario tiene como propósito acreditar los aprendizajes adquiridos al finalizar el curso/taller."),
        ("Indicaciones para el facilitador:", f"La evaluación final se realiza de manera individual y tiene un valor del {data.evaluaciones.pctSumativa or data.evaluaciones.pctSuma or 0}%."),
        ("Condiciones de aplicación:", "Se aplicará al final del curso. El participante no deberá realizar consultas documentales o verbales."),
        ("Tiempo para desarrollar la actividad:", "Considerar el mismo tiempo establecido en el documento de planeación."),
        ("Instrumento:", data.evaluaciones.instSumativa or data.evaluaciones.instSuma or "Cuestionario"),
    ]
    for etiqueta, texto in instrucciones:
        p = doc.add_paragraph()
        p.add_run(etiqueta).bold = True
        p.add_run(f" {texto}")

    doc.add_paragraph()
    doc.add_paragraph("Nombre y firma del Participante: ________________________________________")
    doc.add_paragraph()

    reactivos_texto = (
        data.evaluaciones.instSumativa
        or data.evaluaciones.instSuma
        or ""
    ).strip()

    header_suma = (getattr(data.evaluaciones, "instSumativaHeader", "") or "").strip()
    if header_suma:
        doc.add_paragraph(header_suma)

    if reactivos_texto:
        doc.add_heading("Cuestionario", level=2)
        for linea in reactivos_texto.split("\n"):
            if linea.strip():
                doc.add_paragraph(linea.strip())
    else:
        doc.add_heading("Cuestionario", level=2)
        for i in range(1, 11):
            doc.add_paragraph(f"{i}. __________________________________________________________")
            doc.add_paragraph()

    return _docx_bytes(doc)


def generar_evaluacion_reaccion(data) -> bytes:
    doc = Document()
    _titulo(doc, "EVALUACIÓN DE REACCIÓN")
    tabla_encabezado(doc, data.datos, sin_participante=True)

    doc.add_paragraph("Instrucciones generales: Valore cada aspecto marcando con una X la columna que corresponda.")

    categorias = [
        ("De las características del evento", ["Se encuentra organizado", "Recibí atención cordial"]),
        ("Del contenido del curso", ["Se cubrió el objetivo general", "Se cubrieron los objetivos específicos", "Se realizaron actividades de aprendizaje"]),
        ("De las instalaciones", ["Están limpias", "Iluminación adecuada", "Cuenta con equipo adecuado", "Cuenta con mobiliario adecuado"]),
        ("Del desempeño del instructor", ["Organización de su trabajo", "Claridad al exponer", "Respetó a los participantes", "Promovió la participación"]),
        ("Del material didáctico", ["Estuvo organizado", "Presentaciones claras", "Estuvo entendible", "Está relacionado con el tema"]),
    ]
    escala = ["Excelente 10", "Bueno 8-9", "Regular 6-7", "Malo 5"]

    for categoria, items in categorias:
        doc.add_heading(categoria, level=2)
        t = doc.add_table(rows=len(items) + 1, cols=5)
        t.style = "Table Grid"
        encabezados = ["Aspecto"] + escala
        for j, enc in enumerate(encabezados):
            t.rows[0].cells[j].paragraphs[0].add_run(enc).bold = True
        for i, item in enumerate(items, 1):
            t.rows[i].cells[0].text = item
        doc.add_paragraph()

    doc.add_paragraph("En general, lo que más me gustó del taller fue: _________________________________")
    doc.add_paragraph("En general, lo que menos me gustó del taller fue: _______________________________")
    doc.add_paragraph("Comentarios y sugerencias: _________________________________________________")
    doc.add_paragraph()

    instReac_texto = (data.evaluaciones.instReac or "").strip()
    if instReac_texto:
        doc.add_heading("Preguntas adicionales", level=2)
        for linea in instReac_texto.split("\n"):
            if linea.strip():
                doc.add_paragraph(linea.strip())
        doc.add_paragraph()

    doc.add_paragraph("Gracias por tu participación")

    return _docx_bytes(doc)


def generar_lista_asistencia(data) -> bytes:
    doc = Document()
    _titulo(doc, "LISTA DE ASISTENCIA")
    tabla_encabezado(doc, data.datos, sin_participante=True)

    t = doc.add_table(rows=21, cols=3)
    t.style = "Table Grid"
    encabezados = ["No.", "Nombre", "Firma"]
    for j, enc in enumerate(encabezados):
        t.rows[0].cells[j].paragraphs[0].add_run(enc).bold = True
    for i in range(1, 21):
        t.rows[i].cells[0].text = str(i)

    return _docx_bytes(doc)
