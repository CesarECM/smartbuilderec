import io
import zipfile
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

_VERDE = RGBColor(0x06, 0x5F, 0x46)

_CHECKLIST_ESPACIO = [
    ("esp_materiales",    "Material, herramientas, mobiliario y equipo suficientes para el servicio"),
    ("esp_desplazamiento","Espacio suficiente para desplazamiento libre y distancia con el usuario"),
    ("esp_archivo",       "Archivero/medio digital para resguardo de documentación del usuario"),
    ("esp_ventilacion",   "Ventilado y sin corrientes de aire"),
    ("esp_iluminacion",   "Energía eléctrica e iluminación natural"),
    ("esp_colores",       "Colores claros en paredes y techo"),
    ("esp_basura",        "Depósitos para basura orgánica e inorgánica"),
    ("esp_residuos",      "Depósitos para residuos peligrosos biológico-infecciosos (NOM-087-ECOL-SSA1-2002)"),
    ("esp_pertenencias",  "Lugar específico para que los usuarios coloquen sus pertenencias"),
]

_CHECKLIST_SANITARIO = [
    ("san_manos",      "Lavado de manos conforme al protocolo OMS (antes de recibir al usuario)"),
    ("san_cubrebocas", "Cubrebocas quirúrgico de triple capa colocado correctamente"),
    ("san_tapete",     "Tapete sanitizante con solución desinfectante en ingreso del inmueble"),
    ("san_gel",        "Gel antibacterial disponible en el ingreso"),
    ("san_termometro", "Termómetro digital disponible para toma de temperatura en ingreso"),
]


def _titulo(doc: Document, texto: str, size: int = 14) -> None:
    h = doc.add_heading("", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run(texto)
    r.font.color.rgb = _VERDE
    r.font.size = Pt(size)


def _subtitulo(doc: Document, texto: str) -> None:
    h = doc.add_heading("", level=2)
    r = h.add_run(texto)
    r.font.color.rgb = _VERDE
    r.font.size = Pt(11)


def _linea(doc: Document, campo: str, valor: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(f"{campo}: ")
    r.bold = True
    p.add_run(valor or "—")


def _docx_bytes(doc: Document) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _encabezado_auxiliar(doc: Document, datos: dict, fecha: str = "") -> None:
    nombre = f"{datos.get('aux75Nombre', '')} {datos.get('aux75Apellidos', '')}".strip()
    _linea(doc, "Auxiliar", nombre)
    _linea(doc, "CURP", datos.get("aux75CURP", ""))
    _linea(doc, "Centro / empresa", datos.get("aux75Centro", ""))
    _linea(doc, "Técnicas que aplica", datos.get("aux75Tecnicas", ""))
    _linea(doc, "Fecha de la sesión", fecha or datos.get("aux75Fecha", ""))
    doc.add_paragraph()


# ── Documento 1: Ficha de Registro de Atención Integral ──────────────────────

def generar_ficha_registro(payload: dict) -> bytes:
    doc   = Document()
    datos = payload.get("datos", {})
    us    = payload.get("usuario", {})
    sig   = payload.get("signos", {})
    fecha = datos.get("aux75Fecha", str(date.today()))

    _titulo(doc, "Ficha de Registro de Atención Integral")
    _titulo(doc, "Estándar de Competencia EC1375 — CONOCER", size=11)
    doc.add_paragraph()

    _subtitulo(doc, "Datos del auxiliar")
    _encabezado_auxiliar(doc, datos, fecha)

    _subtitulo(doc, "Datos del usuario")
    _linea(doc, "Nombre completo", us.get("usr75Nombre", ""))
    _linea(doc, "Fecha de nacimiento", us.get("usr75FechaNac", ""))
    _linea(doc, "Edad", us.get("usr75Edad", ""))
    _linea(doc, "Dirección", us.get("usr75Direccion", ""))
    _linea(doc, "Médico tratante", us.get("usr75Medico", ""))
    doc.add_paragraph()

    _subtitulo(doc, "Antecedentes y condición de salud")
    _linea(doc, "Antecedentes físicos/socioemocionales", us.get("usr75Antecedentes", ""))
    _linea(doc, "Enfermedades crónicas, degenerativas y alergias", us.get("usr75Enfermedades", ""))
    _linea(doc, "Hábitos de alimentación, sueño y actividad física", us.get("usr75Habitos", ""))
    _linea(doc, "Motivo / interés para recibir el servicio", us.get("usr75Motivo", ""))
    doc.add_paragraph()

    _subtitulo(doc, "Signos vitales y exploración física (E4324)")
    _linea(doc, "Saturación de oxígeno (SpO₂)", f"{sig.get('sig75SpO2', '')} %")
    _linea(doc, "Pulso", f"{sig.get('sig75Pulso', '')} lpm")
    _linea(doc, "Presión arterial",
           f"{sig.get('sig75PAS', '')}/{sig.get('sig75PAD', '')} mmHg")
    _linea(doc, "Frecuencia respiratoria", f"{sig.get('sig75FrecResp', '')} rpm")
    _linea(doc, "Observaciones de postura", sig.get("sig75Postura", ""))
    doc.add_paragraph()

    doc.add_paragraph(
        "Nota: Los servicios tradicionales y complementarios no cubren ni sustituyen "
        "las indicaciones del médico tratante/profesional de la salud."
    )
    doc.add_paragraph()
    doc.add_paragraph("Firma del auxiliar: ___________________")
    doc.add_paragraph("Firma del usuario:  ___________________")
    return _docx_bytes(doc)


# ── Documento 2: Lista de Verificación del Espacio ───────────────────────────

def generar_lista_verificacion(payload: dict) -> bytes:
    doc  = Document()
    esp  = payload.get("espacio", {})
    datos = payload.get("datos", {})

    _titulo(doc, "Lista de Verificación del Espacio y Protocolos Sanitarios")
    _titulo(doc, "Estándar de Competencia EC1375 — Elemento E4323", size=11)
    doc.add_paragraph()
    _encabezado_auxiliar(doc, datos)

    _subtitulo(doc, "Protocolos de seguridad sanitaria")
    for campo, label in _CHECKLIST_SANITARIO:
        val = esp.get(campo, False)
        p   = doc.add_paragraph(style="List Bullet")
        p.add_run(f"[{'✓' if val else '✗'}]  {label}")

    _subtitulo(doc, "Condiciones del espacio de atención")
    for campo, label in _CHECKLIST_ESPACIO:
        val = esp.get(campo, False)
        p   = doc.add_paragraph(style="List Bullet")
        p.add_run(f"[{'✓' if val else '✗'}]  {label}")

    obs = esp.get("observaciones", "")
    if obs:
        _subtitulo(doc, "Observaciones adicionales")
        doc.add_paragraph(obs)

    doc.add_paragraph()
    doc.add_paragraph("Firma del auxiliar: ___________________")
    return _docx_bytes(doc)


# ── Documento 3: Carta de Consentimiento Informado ───────────────────────────

def generar_consentimiento(payload: dict) -> bytes:
    doc  = Document()
    datos = payload.get("datos", {})
    us   = payload.get("usuario", {})
    con  = payload.get("consentimiento", {})
    fecha = datos.get("aux75Fecha", str(date.today()))
    aux_nombre = f"{datos.get('aux75Nombre', '')} {datos.get('aux75Apellidos', '')}".strip()

    _titulo(doc, "Carta de Consentimiento Informado")
    _titulo(doc, "Aceptación del Servicio Tradicional y Complementario — EC1375", size=11)
    doc.add_paragraph()

    doc.add_paragraph(f"Lugar y fecha: _________________, {fecha}")
    doc.add_paragraph()
    doc.add_paragraph(
        f"Yo, {us.get('usr75Nombre', '_________________')}, hago constar que he recibido "
        f"información completa y comprensible sobre el servicio que se describe a continuación:"
    )
    doc.add_paragraph()

    _subtitulo(doc, "Servicio a realizar")
    _linea(doc, "Técnica designada", con.get("con75Tecnica", ""))
    _linea(doc, "Descripción", con.get("con75Descripcion", ""))
    _linea(doc, "Objetivo de la sesión", con.get("con75Objetivo", ""))
    _linea(doc, "Reacciones / sensaciones posibles", con.get("con75Reacciones", ""))
    _linea(doc, "Vestimenta recomendada", con.get("con75Vestimenta", ""))
    _linea(doc, "Número de sesiones", con.get("con75NumSesiones", ""))
    _linea(doc, "Duración por sesión", f"{con.get('con75Duracion', '')} minutos")
    doc.add_paragraph()

    _subtitulo(doc, "Declaración del usuario")
    doc.add_paragraph(
        "Declaro que los datos personales proporcionados son verídicos y que he sido informado(a) "
        "de que los servicios tradicionales y complementarios no cubren ni sustituyen las "
        "indicaciones del médico tratante/profesional de la salud, de conformidad con la "
        "Ley Federal de Protección de Datos Personales en Posesión de Particulares."
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "Manifiesto mi consentimiento libre e informado para recibir el servicio descrito."
    )
    doc.add_paragraph()
    doc.add_paragraph("Firma del usuario:  ___________________")
    doc.add_paragraph(f"Nombre: {us.get('usr75Nombre', '')}")
    doc.add_paragraph()
    doc.add_paragraph("Firma del auxiliar: ___________________")
    doc.add_paragraph(f"Nombre: {aux_nombre}")
    doc.add_paragraph(f"Centro: {datos.get('aux75Centro', '')}")
    return _docx_bytes(doc)


# ── Documento 4: Plan de Seguimiento y Sesiones ───────────────────────────────

def generar_plan_seguimiento(payload: dict) -> bytes:
    doc  = Document()
    datos = payload.get("datos", {})
    us   = payload.get("usuario", {})
    seg  = payload.get("seguimiento", {})
    con  = payload.get("consentimiento", {})
    fecha = datos.get("aux75Fecha", str(date.today()))
    aux_nombre = f"{datos.get('aux75Nombre', '')} {datos.get('aux75Apellidos', '')}".strip()

    _titulo(doc, "Plan de Seguimiento y Sesiones")
    _titulo(doc, "Estándar de Competencia EC1375 — Elemento E4326", size=11)
    doc.add_paragraph()

    _subtitulo(doc, "Datos generales")
    _linea(doc, "Usuario", us.get("usr75Nombre", ""))
    _linea(doc, "Auxiliar", aux_nombre)
    _linea(doc, "Centro", datos.get("aux75Centro", ""))
    _linea(doc, "Fecha de inicio", fecha)
    _linea(doc, "Técnica aplicada", con.get("con75Tecnica", ""))
    _linea(doc, "Motivo del servicio", us.get("usr75Motivo", ""))
    doc.add_paragraph()

    _subtitulo(doc, "Programa de seguimiento")
    _linea(doc, "Número de sesiones programadas", seg.get("seg75NumSesiones", ""))
    _linea(doc, "Frecuencia", seg.get("seg75Frecuencia", ""))
    _linea(doc, "Duración por sesión", f"{seg.get('seg75Duracion', '')} minutos")
    _linea(doc, "Vía de contacto",     seg.get("seg75ContactoVia", ""))
    _linea(doc, "Datos de contacto",   seg.get("seg75Contacto", ""))
    doc.add_paragraph()

    if seg.get("seg75ObjetivoSesion"):
        _subtitulo(doc, "Objetivos por sesión")
        doc.add_paragraph(seg["seg75ObjetivoSesion"])
        doc.add_paragraph()

    if seg.get("seg75TareasEnCasa"):
        _subtitulo(doc, "Actividades / ejercicios en casa")
        doc.add_paragraph(seg["seg75TareasEnCasa"])
        doc.add_paragraph()

    if seg.get("seg75Recomendaciones"):
        _subtitulo(doc, "Recomendaciones del especialista")
        doc.add_paragraph(seg["seg75Recomendaciones"])
        doc.add_paragraph()

    _subtitulo(doc, "Plan de sesiones")
    nsesiones = int(seg.get("seg75NumSesiones", 0) or 0)
    if nsesiones > 0:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, txt in enumerate(["Sesión", "Fecha", "Hora inicio/fin", "Actividades / intervención", "Notas de evolución"]):
            hdr[i].text = txt
        for n in range(1, nsesiones + 1):
            row = table.add_row().cells
            row[0].text = str(n)
    doc.add_paragraph()

    doc.add_paragraph("Firma del auxiliar: ___________________")
    doc.add_paragraph(f"Nombre: {aux_nombre}")
    doc.add_paragraph()
    doc.add_paragraph("Firma del usuario:  ___________________")
    doc.add_paragraph(f"Nombre: {us.get('usr75Nombre', '')}")
    return _docx_bytes(doc)


# ── ZIP principal ─────────────────────────────────────────────────────────────

def generar_zip_ec1375(payload: dict) -> bytes:
    docs = [
        ("01_Ficha_Registro_Atencion.docx",        generar_ficha_registro(payload)),
        ("02_Lista_Verificacion_Espacio.docx",      generar_lista_verificacion(payload)),
        ("03_Consentimiento_Informado.docx",        generar_consentimiento(payload)),
        ("04_Plan_Seguimiento_Sesiones.docx",       generar_plan_seguimiento(payload)),
    ]
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
        for nombre, contenido in docs:
            zf.writestr(nombre, contenido)
    buf.seek(0)
    return buf.read()
