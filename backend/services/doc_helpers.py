import io
import json
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = Path(__file__).resolve().parent.parent


def load_prompt(filename: str, **kwargs) -> str:
    file_path = BASE_DIR / "prompts" / filename
    if not file_path.exists():
        raise FileNotFoundError(f"No se encontró el archivo {filename}")
    with open(file_path, "r", encoding="utf-8") as f:
        content = f.read()
    return content.format(**kwargs)


def tabla_encabezado(doc: Document, datos) -> None:
    """Agrega la tabla de encabezado estándar EC0217 al documento."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tabla = doc.add_table(rows=4, cols=2)
    tabla.style = "Table Grid"

    def set_cell(row, col, label, value=""):
        c = tabla.rows[row].cells[col]
        c.text = ""
        p = c.paragraphs[0]
        run_label = p.add_run(label)
        run_label.bold = True
        if value:
            p.add_run(f" {value}")

    set_cell(0, 0, "Nombre del Curso/sesión:", datos.nombreCurso)
    set_cell(0, 1, "Lugar de Instrucción:", datos.lugar)
    set_cell(1, 0, "Nombre del facilitador/instructor/capacitador/formador:", datos.instructor)
    set_cell(1, 1, "Nombre del Diseñador:", datos.disenador)
    set_cell(2, 0, "Duración del curso:", f"{datos.duracion} minutos")
    set_cell(2, 1, "Fecha(s):", datos.fecha)
    tabla.rows[3].cells[0].merge(tabla.rows[3].cells[1])
    set_cell(3, 0, "Nombre del Participante:", "")

    doc.add_paragraph()


def _docx_bytes(doc: Document) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _titulo(doc: Document, texto: str, subtitulo: str = "") -> None:
    h = doc.add_heading("", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run(texto)
    r.font.color.rgb = RGBColor(0x1F, 0x3B, 0x6D)
    r.font.size = Pt(16)
    if subtitulo:
        h2 = doc.add_heading("", level=2)
        h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = h2.add_run(subtitulo)
        r2.font.color.rgb = RGBColor(0x31, 0x4E, 0x7A)
        r2.font.size = Pt(13)


def crear_zip_con_docx(docx_bytes: bytes, nombre: str = "objetivos_EC0217.docx") -> bytes:
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(nombre, docx_bytes)
    zip_buffer.seek(0)
    return zip_buffer.read()


def limpiar_nombre_archivo(texto: str) -> str:
    texto = str(texto or "").strip()
    texto = re.sub(r"\s+", "_", texto)
    texto = re.sub(r'[\\/:*?"<>|]', "", texto)
    texto = re.sub(r"_+", "_", texto)
    return texto or "Archivo"


def _validar_expediente(data) -> None:
    errores = []

    if not (data.datos.nombreCurso or "").strip():
        errores.append("Nombre del curso es obligatorio (Paso 1).")
    if not (data.datos.instructor or "").strip():
        errores.append("Nombre del instructor es obligatorio (Paso 1).")
    if not data.datos.duracion or data.datos.duracion < 120:
        errores.append("La duración debe ser mínimo 120 minutos según la norma EC0217 (Paso 1).")
    if not data.datos.participantes or data.datos.participantes < 1:
        errores.append("El número de participantes es obligatorio (Paso 1).")

    pct_diag = data.evaluaciones.pctDiagnostica or data.evaluaciones.pctDiag or 0
    pct_form = data.evaluaciones.pctFormativa  or data.evaluaciones.pctForm or 0
    pct_suma = data.evaluaciones.pctSumativa   or data.evaluaciones.pctSuma or 0
    total_pct = pct_diag + pct_form + pct_suma
    if total_pct != 100:
        errores.append(f"Los porcentajes de evaluación deben sumar 100% (actualmente suman {total_pct}%) (Paso 14).")

    temas_totales = len(data.temario.u1) + len(data.temario.u2) + len(data.temario.u3)
    if temas_totales == 0:
        errores.append("El temario debe tener al menos un tema en alguna unidad (Paso 4).")

    obj = data.objetivos
    if not any([obj.general, obj.cognitiva, obj.psicomotriz, obj.afectiva]):
        errores.append("Debes completar al menos un objetivo de aprendizaje (Paso 2).")

    if errores:
        from fastapi import HTTPException
        raise HTTPException(status_code=422, detail=errores)


def _calc_content_len(data) -> int:
    """Longitud total de campos clave para comparar versiones del mismo curso."""
    try:
        parts = [
            data.datos.instructor or "",
            data.datos.nombreCurso or "",
            data.beneficios or "",
            json.dumps(data.objetivos.model_dump()),
            json.dumps(data.temario.model_dump()),
            json.dumps(data.expositiva or {}),
            json.dumps(data.demostrativa or {}),
            json.dumps(data.dialogo or {}),
        ]
        return sum(len(p) for p in parts)
    except Exception:
        return 0
