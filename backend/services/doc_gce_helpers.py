import io
import zipfile

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

_VERDE = RGBColor(0x06, 0x5F, 0x46)

_ESCALA_ENC = {1: "Muy insatisfecho", 2: "Insatisfecho", 3: "Satisfecho", 4: "Muy satisfecho"}

_PREGUNTAS_ENC = [
    "¿La evaluación se realizó conforme al Plan de Evaluación?",
    "¿El evaluador le explicó claramente el proceso de evaluación?",
    "¿Las instalaciones y condiciones fueron adecuadas?",
    "¿Los instrumentos de evaluación fueron claros y comprensibles?",
    "¿El tiempo asignado para la evaluación fue suficiente?",
    "¿El evaluador le trató con respeto y profesionalismo?",
    "¿Recibió información clara sobre el resultado de su evaluación?",
    "En general, ¿quedó satisfecho con el proceso de evaluación?",
]


# ── Helpers base ──────────────────────────────────────────────────────────────

def _titulo(doc: Document, texto: str, size: int = 13) -> None:
    h = doc.add_heading("", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run(texto)
    r.font.color.rgb = _VERDE
    r.font.size = Pt(size)


def _sub(doc: Document, texto: str) -> None:
    h = doc.add_heading(texto, level=2)
    for r in h.runs:
        r.font.color.rgb = _VERDE
        r.font.size = Pt(11)


def _campo(doc: Document, label: str, valor: str) -> None:
    p = doc.add_paragraph()
    p.add_run(f"{label}: ").bold = True
    p.add_run(str(valor or "—"))


def _firma(doc: Document, label: str, nombre: str = "") -> None:
    doc.add_paragraph(f"{label}: {nombre or '___________________________'}")


def _docx_bytes(doc: Document) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _cab(doc: Document, titulo_doc: str, datos: dict, ec: dict) -> None:
    _titulo(doc, titulo_doc)
    _titulo(doc, f"{ec.get('codigo', '')} — {ec.get('titulo', '')}", size=10)
    doc.add_paragraph()
    f = datos.get("ficha_registro", {})
    _campo(doc, "Candidato", f.get("nombre_completo", ""))
    _campo(doc, "CURP", f.get("curp", ""))
    doc.add_paragraph()


# ── Doc 1: Ficha de Registro ──────────────────────────────────────────────────

def _generar_ficha(datos: dict, ec: dict) -> bytes:
    doc = Document()
    _cab(doc, "Ficha de Registro del Candidato", datos, ec)
    f = datos.get("ficha_registro", {})

    _sub(doc, "Datos personales")
    _campo(doc, "Fecha de nacimiento", f.get("fecha_nacimiento", ""))
    _campo(doc, "Género", f.get("genero", ""))
    _campo(doc, "Lugar de nacimiento", f.get("lugar_nacimiento", ""))
    _campo(doc, "Nacionalidad", f.get("nacionalidad", ""))
    _campo(doc, "Correo electrónico", f.get("email", ""))
    _campo(doc, "Teléfono", f.get("telefono", ""))
    dom = f.get("domicilio", {})
    if isinstance(dom, dict):
        dom_str = " ".join(filter(None, [
            dom.get("calle", ""), dom.get("num_ext", ""), dom.get("colonia", ""),
            dom.get("ciudad", ""), dom.get("estado", ""),
            f"C.P. {dom.get('cp', '')}" if dom.get("cp") else "",
        ]))
        _campo(doc, "Domicilio", dom_str)
    else:
        _campo(doc, "Domicilio", str(dom or ""))

    _sub(doc, "Escolaridad y experiencia")
    _campo(doc, "Sabe leer y escribir", "Sí" if f.get("sabe_leer_escribir") else "No")
    _campo(doc, "Último grado de estudios", f.get("estudios", ""))
    _campo(doc, "Discapacidad", f.get("discapacidad", "Ninguna"))
    _campo(doc, "Trabaja actualmente", "Sí" if f.get("trabaja_actualmente") else "No")
    _campo(doc, "Puesto actual", f.get("puesto", ""))
    _campo(doc, "Experiencia en el área", f.get("experiencia", ""))
    _campo(doc, "Certificaciones previas", f.get("certificaciones_previas", ""))

    _sub(doc, "Consentimiento")
    _campo(doc, "Acepta uso de datos (RENAP)", "Sí" if f.get("consentimiento_renap") else "No")
    _campo(doc, "Observaciones", f.get("observaciones", ""))
    doc.add_paragraph()
    _firma(doc, "Firma del candidato")
    return _docx_bytes(doc)


# ── Doc 2: Diagnóstico ────────────────────────────────────────────────────────

def _generar_diagnostico(datos: dict, ec: dict) -> bytes:
    doc = Document()
    _cab(doc, "Diagnóstico Inicial", datos, ec)
    d   = datos.get("diagnostico", {})
    cfg = ec.get("config", {}).get("diagnostico", {})

    _sub(doc, "Resultado")
    total = d.get("total") or cfg.get("total_reactivos", 0)
    _campo(doc, "Respuestas correctas", f"{d.get('correctas', '—')} / {total}")
    _campo(doc, "Posibilidad de éxito", d.get("posibilidad", "—"))
    _campo(doc, "Sugerencia", d.get("sugerencia", "—"))
    _campo(doc, "Decisión del candidato", d.get("decision_candidato", "—"))
    _campo(doc, "Fecha", d.get("fecha", "—"))
    doc.add_paragraph()

    tabla = cfg.get("tabla_interpretacion", [])
    if tabla:
        _sub(doc, "Tabla de interpretación")
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"
        hd = tbl.rows[0].cells
        hd[0].text = "Mín."; hd[1].text = "Máx."; hd[2].text = "Posibilidad"
        for row in tabla:
            c = tbl.add_row().cells
            c[0].text = str(row.get("min", ""))
            c[1].text = str(row.get("max", ""))
            c[2].text = row.get("posibilidad", "")

    doc.add_paragraph()
    _firma(doc, "Firma del candidato")
    _firma(doc, "Firma del evaluador")
    return _docx_bytes(doc)


# ── Doc 3: Plan de Evaluación ─────────────────────────────────────────────────

def _generar_plan(datos: dict, ec: dict) -> bytes:
    doc = Document()
    _cab(doc, "Plan de Evaluación", datos, ec)
    p   = datos.get("plan_evaluacion", {})
    cfg = ec.get("config", {}).get("plan_evaluacion", {})

    _sub(doc, "Datos logísticos")
    _campo(doc, "Fecha de evaluación", p.get("fecha_evaluacion", ""))
    _campo(doc, "Horario", p.get("horario_evaluacion", ""))
    _campo(doc, "Lugar de evaluación", p.get("lugar_evaluacion", ""))
    _campo(doc, "Fecha de resultados", p.get("fecha_resultados", ""))
    _campo(doc, "Horario resultados", p.get("horario_resultados", ""))
    _campo(doc, "Lugar de resultados", p.get("lugar_resultados", ""))
    doc.add_paragraph()

    acts = cfg.get("actividades", [])
    if acts:
        _sub(doc, "Actividades del plan")
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"
        hd = tbl.rows[0].cells
        hd[0].text = "#"; hd[1].text = "Técnica / Instrumento"; hd[2].text = "Descripción"
        for a in acts:
            c = tbl.add_row().cells
            c[0].text = str(a.get("num", ""))
            c[1].text = a.get("tecnica_instrumento", a.get("tecnica", ""))
            c[2].text = a.get("descripcion", "")

    if cfg.get("puntaje_minimo") or cfg.get("criterio_1"):
        doc.add_paragraph()
        _sub(doc, "Criterios para el juicio")
        if cfg.get("puntaje_minimo"):
            _campo(doc, "Puntaje mínimo", str(cfg["puntaje_minimo"]))
        if cfg.get("criterio_1"):
            _campo(doc, "Criterio 1", cfg["criterio_1"])
        if cfg.get("criterio_2"):
            _campo(doc, "Criterio 2", cfg["criterio_2"])

    doc.add_paragraph()
    _firma(doc, "Evaluador", p.get("firma_evaluador", ""))
    _firma(doc, "Candidato", p.get("firma_candidato", ""))
    return _docx_bytes(doc)


# ── Doc 4: IEC Aplicado ───────────────────────────────────────────────────────

def _generar_iec(datos: dict, ec: dict) -> bytes:
    doc = Document()
    _cab(doc, "Instrumento de Evaluación de la Competencia (IEC)", datos, ec)
    d         = datos.get("iec", {})
    cfg_react = ec.get("config", {}).get("iec", {}).get("reactivos", [])
    react_map = {r["codigo"]: r for r in (d.get("reactivos") or [])}

    _campo(doc, "Fecha de aplicación", d.get("fecha_aplicacion", ""))
    doc.add_paragraph()

    if cfg_react:
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Table Grid"
        hd = tbl.rows[0].cells
        hd[0].text = "Código"; hd[1].text = "Descripción"
        hd[2].text = "C / NC"; hd[3].text = "Peso %"
        for r in cfg_react:
            rd     = react_map.get(r["codigo"], {})
            cumple = rd.get("cumple")
            c = tbl.add_row().cells
            c[0].text = r.get("codigo", "")
            c[1].text = r.get("descripcion", "")
            c[2].text = "C" if cumple is True else ("NC" if cumple is False else "—")
            c[3].text = str(r.get("peso_relativo", ""))

    doc.add_paragraph()
    _campo(doc, "Puntaje obtenido", d.get("puntaje_obtenido", "—"))
    _campo(doc, "Juicio preliminar",
           "COMPETENTE" if d.get("juicio_preliminar") == "C"
           else ("TODAVÍA NO COMPETENTE" if d.get("juicio_preliminar") == "TNC" else "—"))
    doc.add_paragraph()
    _firma(doc, "Evaluador")
    _firma(doc, "Candidato")
    return _docx_bytes(doc)


# ── Doc 5: Cédula de Evaluación ───────────────────────────────────────────────

def _generar_cedula(datos: dict, ec: dict) -> bytes:
    doc = Document()
    _cab(doc, "Cédula de Evaluación", datos, ec)
    ced    = datos.get("cedula", {})
    juicio = ced.get("juicio", "")

    _campo(doc, "Fecha", ced.get("fecha", ""))
    _campo(doc, "Juicio de evaluación",
           "COMPETENTE" if juicio == "C"
           else ("TODAVÍA NO COMPETENTE" if juicio == "TNC" else juicio or "—"))
    doc.add_paragraph()
    _sub(doc, "Mejores prácticas observadas")
    doc.add_paragraph(ced.get("mejores_practicas", "") or "—")
    _sub(doc, "Áreas de oportunidad")
    doc.add_paragraph(ced.get("areas_oportunidad", "") or "—")
    _sub(doc, "Recomendaciones")
    doc.add_paragraph(ced.get("recomendaciones", "") or "—")
    doc.add_paragraph()
    _firma(doc, "Evaluador")
    _firma(doc, "Candidato")
    return _docx_bytes(doc)


# ── Doc 6: Encuesta de Satisfacción ──────────────────────────────────────────

def _generar_encuesta(datos: dict, ec: dict) -> bytes:
    doc = Document()
    _cab(doc, "Encuesta de Satisfacción del Candidato", datos, ec)
    enc  = datos.get("encuesta", {})
    resp = enc.get("respuestas", [])

    _campo(doc, "Fecha", enc.get("fecha", ""))
    doc.add_paragraph()

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    hd = tbl.rows[0].cells
    hd[0].text = "#"; hd[1].text = "Pregunta"; hd[2].text = "Calificación"
    for i, preg in enumerate(_PREGUNTAS_ENC):
        val   = resp[i] if i < len(resp) and resp[i] is not None else None
        label = f"{val} — {_ESCALA_ENC.get(val, '')}" if val else "—"
        c = tbl.add_row().cells
        c[0].text = str(i + 1)
        c[1].text = preg
        c[2].text = label

    doc.add_paragraph()
    _sub(doc, "Observaciones adicionales")
    doc.add_paragraph(enc.get("observaciones", "") or "—")
    doc.add_paragraph()
    _firma(doc, "Candidato")
    return _docx_bytes(doc)


# ── ZIP ───────────────────────────────────────────────────────────────────────

def generar_zip_portafolio(datos: dict, ec: dict) -> bytes:
    documentos = [
        ("01_Ficha_de_Registro.docx",         _generar_ficha(datos, ec)),
        ("02_Diagnostico.docx",               _generar_diagnostico(datos, ec)),
        ("03_Plan_de_Evaluacion.docx",        _generar_plan(datos, ec)),
        ("04_IEC_Aplicado.docx",              _generar_iec(datos, ec)),
        ("05_Cedula_de_Evaluacion.docx",      _generar_cedula(datos, ec)),
        ("06_Encuesta_de_Satisfaccion.docx",  _generar_encuesta(datos, ec)),
    ]
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
        for nombre, contenido in documentos:
            zf.writestr(nombre, contenido)
    buf.seek(0)
    return buf.read()
