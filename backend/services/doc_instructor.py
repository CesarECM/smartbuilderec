import json
import os
from concurrent.futures import ThreadPoolExecutor

from docx import Document
from docx.shared import Pt, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openai import OpenAI

from services.doc_helpers import load_prompt, _docx_bytes, _titulo
from services.doc_planeacion_docx import _preprocesar
from services.doc_instructor_secciones import (
    _seccion_carta_descriptiva,
    _seccion_requerimientos,
    _seccion_instrumentos,
    _seccion_unidad_instructor,
)

_client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
_MODEL  = os.getenv("OPENAI_MODEL_GRL", "gpt-4o-mini")


def _ia_introduccion(data) -> str:
    datos = data.datos
    obj   = data.objetivos
    tem   = data.temario
    exp   = data.expositiva or {}
    dem   = data.demostrativa or {}
    dia   = data.dialogo or {}

    def _lista(items): return ", ".join(i for i in (items or []) if i) or "—"

    prompt = load_prompt(
        "manual_instructor_intro_prompt.txt",
        nombreCurso=datos.nombreCurso or "",
        objetivoGeneral=obj.general or "",
        tecnicaExpositiva=exp.get("tema") or exp.get("descripcion") or "Técnica expositiva",
        tecnicaDemostrativa=dem.get("actividad") or "Técnica demostrativa",
        tecnicaDialogo=dia.get("instrucciones") or "Diálogo y discusión grupal",
        temasU1=_lista(tem.u1),
        temasU2=_lista(tem.u2),
        temasU3=_lista(tem.u3),
    )
    resp = _client.chat.completions.create(
        model=_MODEL,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.4,
        response_format={"type": "json_object"},
    )
    return json.loads(resp.choices[0].message.content.strip()).get("introduccion", "")


def _ia_sugerencias_unidad(data, num_unidad: int, temas: list, obj_particular: str) -> list:
    datos = data.datos
    obj   = data.objetivos
    exp   = data.expositiva or {}
    dem   = data.demostrativa or {}
    dia   = data.dialogo or {}

    tecnica = (
        f"Expositiva ({exp.get('tema') or 'presentación de contenido al grupo'}), "
        f"Demostrativa ({dem.get('actividad') or 'demostración y práctica guiada'}), "
        f"Diálogo/Discusión ({(dia.get('instrucciones') or 'reflexión y análisis grupal')[:80]})"
    )
    prompt = load_prompt(
        "manual_instructor_sugerencias_prompt.txt",
        nombreCurso=datos.nombreCurso or "",
        objetivoGeneral=obj.general or "",
        numUnidad=str(num_unidad),
        objetivoParticular=obj_particular or "",
        temas="\n".join(f"- {t}" for t in temas) if temas else "—",
        tecnicaInstruccional=tecnica,
    )
    resp = _client.chat.completions.create(
        model=_MODEL,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.4,
        response_format={"type": "json_object"},
    )
    return json.loads(resp.choices[0].message.content.strip()).get("sugerencias", [])


def generar_manual_instructor(data) -> bytes:
    payload = data.model_dump()
    d       = _preprocesar(payload)

    obj = data.objetivos
    tem = data.temario
    obj_particulares = [obj.cognitiva or "", obj.psicomotriz or "", obj.afectiva or ""]
    unidades = [
        (1, tem.u1 or [], obj_particulares[0]),
        (2, tem.u2 or [], obj_particulares[1]),
        (3, tem.u3 or [], obj_particulares[2]),
    ]

    # ── Llamadas IA en paralelo ────────────────────────────────────────────────
    intro_texto          = ""
    sugerencias_unidades = {1: [], 2: [], 3: []}

    with ThreadPoolExecutor(max_workers=4) as exc:
        fut_intro = exc.submit(_ia_introduccion, data)
        fut_sugs  = {
            num: exc.submit(_ia_sugerencias_unidad, data, num, temas, obj_part)
            for num, temas, obj_part in unidades
        }
        try:
            intro_texto = fut_intro.result()
        except Exception as e:
            print(f"[instructor] IA intro error: {e}")
        for num, _, _ in unidades:
            try:
                sugerencias_unidades[num] = fut_sugs[num].result()
            except Exception as e:
                print(f"[instructor] IA sugerencias U{num} error: {e}")

    # ── Construcción del documento ─────────────────────────────────────────────
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Twips(12240); sec.page_height = Twips(15840)
    sec.left_margin = sec.right_margin = Twips(720)
    sec.top_margin  = sec.bottom_margin = Twips(720)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)

    datos = data.datos

    # Portada
    _titulo(doc, "MANUAL DEL INSTRUCTOR", subtitulo=datos.nombreCurso or "")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disenador = datos.disenador or datos.instructor or ""
    p.add_run(f"Diseñador del curso: {disenador}").bold = True
    doc.add_paragraph()

    # Índice
    doc.add_heading("ÍNDICE", level=1)
    doc.add_paragraph(
        "[El índice se genera automáticamente en Word: Referencias → Tabla de contenido]"
    )
    doc.add_paragraph()

    # Introducción (IA)
    doc.add_heading("INTRODUCCIÓN", level=1)
    doc.add_paragraph(intro_texto or "[Completar la introducción del manual del instructor]")
    doc.add_paragraph()

    # Carta descriptiva completa (reutiliza helpers de planeación)
    _seccion_carta_descriptiva(doc, d)

    # Requerimientos del lugar de capacitación
    _seccion_requerimientos(doc, data)

    # Sugerencias didácticas por unidad (IA)
    for num, temas, obj_part in unidades:
        _seccion_unidad_instructor(doc, num, temas, obj_part, sugerencias_unidades[num])

    # Instrumentos de evaluación con claves
    _seccion_instrumentos(doc, data.evaluaciones)

    # Fuentes de información
    doc.add_heading("FUENTES DE INFORMACIÓN", level=1)
    refs = (data.cierre or {}).get("referencias", "")
    doc.add_paragraph(
        refs if refs else "[Incluir las fuentes de información documental y/o de internet del curso]"
    )

    return _docx_bytes(doc)
