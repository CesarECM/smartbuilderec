import io

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def crear_docx_objetivos(data) -> bytes:
    doc = Document()
    titulo = doc.add_heading("", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_titulo = titulo.add_run("Objetivos de Aprendizaje")
    run_titulo.font.color.rgb = RGBColor(0x1F, 0x3B, 0x6D)
    run_titulo.font.size = Pt(18)
    doc.add_paragraph()
    secciones = [
        ("Objetivo General",     data.general),
        ("Objetivo Cognitivo",   data.cognitiva),
        ("Objetivo Psicomotriz", data.psicomotriz),
        ("Objetivo Afectivo",    data.afectiva),
    ]
    for titulo_sec, texto in secciones:
        h = doc.add_heading("", level=2)
        run_h = h.add_run(titulo_sec)
        run_h.font.color.rgb = RGBColor(0x31, 0x4E, 0x7A)
        run_h.font.size = Pt(14)
        p = doc.add_paragraph()
        run_p = p.add_run(texto)
        run_p.font.size = Pt(12)
        doc.add_paragraph()
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
