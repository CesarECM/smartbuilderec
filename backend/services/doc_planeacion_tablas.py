from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Twips

from services.doc_planeacion_xml import (
    TW, AZUL, AZUL_CLARO, GRIS_CLARO, AMARILLO, NEGRO,
    _table_setup, _hdr, _subhdr, _lbl, _val, _val_s,
    _borders, _pad, _cw, _shading, _clear, _para, _run, _lines,
    _fmt_dur, _inst,
)


def _tabla_info_general(doc, d):
    C1 = round(TW * .20); C2 = round(TW * .30)
    C3 = round(TW * .18); C4 = TW - C1 - C2 - C3
    dat = d.get('datos') or {}

    t = doc.add_table(rows=0, cols=4)
    _table_setup(t, [C1, C2, C3, C4])

    r = t.add_row(); r.cells[0].merge(r.cells[3])
    _hdr(r.cells[0], "INFORMACIÓN GENERAL", TW)

    r = t.add_row(); r.cells[1].merge(r.cells[3])
    _lbl(r.cells[0], "Nombre del Curso / Sesión:", C1)
    _borders(r.cells[1]); _pad(r.cells[1]); _cw(r.cells[1], C2+C3+C4); _clear(r.cells[1])
    _para(r.cells[1], dat.get('nombreCurso',''), bold=True, size_pt=12, sp_before=20, sp_after=20)

    for lbl_txt, key in [("Nombre del Diseñador:", 'disenador'),
                          ("Nombre del Instructor / Facilitador:", 'instructor')]:
        r = t.add_row(); r.cells[1].merge(r.cells[3])
        _lbl(r.cells[0], lbl_txt, C1)
        _val(r.cells[1], dat.get(key, ''), C2+C3+C4)

    r = t.add_row()
    _lbl(r.cells[0], "Lugar de Instrucción:", C1); _val(r.cells[1], dat.get('lugar',''), C2)
    _lbl(r.cells[2], "Duración:", C3);            _val(r.cells[3], _fmt_dur(dat.get('duracion',0)), C4)

    r = t.add_row()
    _lbl(r.cells[0], "Fecha(s):", C1);            _val(r.cells[1], dat.get('fecha',''), C2)
    _lbl(r.cells[2], "Nº de participantes:", C3)
    part = dat.get('participantes')
    _val(r.cells[3], f"{part} participantes" if part else '', C4)

    r = t.add_row(); r.cells[1].merge(r.cells[3])
    _lbl(r.cells[0], "Perfil del participante:", C1)
    _val(r.cells[1], dat.get('perfil',''), C2+C3+C4)

    r = t.add_row(); r.cells[1].merge(r.cells[3])
    _lbl(r.cells[0], "Beneficios del curso:", C1)
    _val(r.cells[1], d.get('beneficios',''), C2+C3+C4, shade=AMARILLO)


def _tabla_objetivos(doc, d):
    C1 = round(TW * .65); C2 = TW - C1
    obj = d.get('objetivos') or {}; tem = d.get('temario') or {}

    t = doc.add_table(rows=0, cols=2)
    _table_setup(t, [C1, C2])

    DESCRIP_OG = (
        "describe la demostración de un conocimiento, desempeño o producto, resultado del "
        "aprendizaje del participante, así como el dominio de aprendizaje cognitivo, "
        "psicomotriz, afectivo y relacional-social en los que impactará el Curso/sesión"
    )

    r = t.add_row(); r.cells[0].merge(r.cells[1])
    _borders(r.cells[0]); _pad(r.cells[0]); _shading(r.cells[0], AZUL_CLARO); _clear(r.cells[0])
    p = r.cells[0].add_paragraph()
    p.paragraph_format.space_before = Twips(20); p.paragraph_format.space_after = Twips(20)
    _run(p, "Objetivo General ", bold=True, color=AZUL)
    _run(p, f"({DESCRIP_OG})", color=NEGRO, size_pt=9)

    r = t.add_row(); r.cells[0].merge(r.cells[1])
    _borders(r.cells[0]); _pad(r.cells[0]); _clear(r.cells[0])
    _para(r.cells[0], obj.get('general',''), sp_before=20, sp_after=20)

    r = t.add_row()
    _subhdr(r.cells[0], "Objetivos Particulares", C1)
    _subhdr(r.cells[1], "Temas:", C2)

    for num, tipo, key_obj, key_tem in [
        ("1", "(Cognitivo)",                   'cognitiva',   'u1'),
        ("2", "(Psicomotriz)",                  'psicomotriz', 'u2'),
        ("3", "(Afectivo / Relacional-social)", 'afectiva',    'u3'),
    ]:
        r = t.add_row()
        _borders(r.cells[0]); _pad(r.cells[0]); _cw(r.cells[0], C1); _clear(r.cells[0])
        p = r.cells[0].add_paragraph()
        p.paragraph_format.space_before = Twips(16); p.paragraph_format.space_after = Twips(8)
        _run(p, f"{num}. ", bold=True); _run(p, tipo, bold=True, color=AZUL)
        _para(r.cells[0], obj.get(key_obj,''), sp_before=8, sp_after=16)
        _val(r.cells[1], '\n'.join(tem.get(key_tem) or []), C2)


def _tabla_requerimientos(doc, d):
    CW = round(TW / 6); CW6 = TW - CW * 5
    mat = d.get('materiales') or {}
    vals = [
        mat.get('instalaciones') or mat.get('integracion')      or "Aula iluminada, ventilada, mesas y sillas suficientes.",
        mat.get('equipo')        or mat.get('expositiva')        or "Laptop, proyector, extensión eléctrica y presentación.",
        mat.get('materialesDidacticos') or mat.get('demostrativa') or "Manual del participante, hojas, bolígrafos y materiales de apoyo.",
        mat.get('humanos')       or mat.get('dialogo')           or "Instructor y participantes registrados.",
        mat.get('otros')         or mat.get('energizante')       or "Material extra requerido para las actividades.",
        mat.get('seguridad') or "Salidas de emergencia señalizadas, extintor vigente, botiquín de primeros auxilios, medidas de higiene y protocolo de evacuación.",
    ]
    labels = [
        "Instalaciones, mobiliario y distribución:",
        "Equipo de apoyo y distribución:",
        "Materiales didácticos de apoyo:",
        "Requerimientos humanos:",
        "Otros:",
        "Salud / Seguridad / Higiene / Protección civil:",
    ]
    widths = [CW, CW, CW, CW, CW, CW6]

    t = doc.add_table(rows=0, cols=6)
    _table_setup(t, widths)

    r = t.add_row(); r.cells[0].merge(r.cells[5])
    _hdr(r.cells[0], "REQUERIMIENTOS PARA EL DESARROLLO DEL CURSO", TW)

    r = t.add_row()
    for i, (lbl, w) in enumerate(zip(labels, widths)):
        _borders(r.cells[i]); _pad(r.cells[i], 60, 60, 80, 80)
        _shading(r.cells[i], AZUL_CLARO); _cw(r.cells[i], w); _clear(r.cells[i])
        _para(r.cells[i], lbl, bold=True, color=AZUL, size_pt=9, sp_before=16, sp_after=16)

    r = t.add_row()
    for i, (val, w) in enumerate(zip(vals, widths)):
        _borders(r.cells[i]); _pad(r.cells[i], 60, 60, 80, 80); _cw(r.cells[i], w)
        _lines(r.cells[i], val, size_pt=9)


def _tabla_evaluaciones(doc, d):
    ev = d.get('evaluaciones') or {}
    C1 = round(TW * .35); C2 = round(TW * .08)
    C3 = round(TW * .37); C4 = TW - C1 - C2 - C3
    pct_d = ev.get('pctDiagnostica') or ev.get('pctDiag') or 0
    pct_f = ev.get('pctFormativa')   or ev.get('pctForm') or 0
    pct_s = ev.get('pctSumativa')    or ev.get('pctSuma') or 0

    t = doc.add_table(rows=0, cols=4)
    _table_setup(t, [C1, C2, C3, C4])

    r = t.add_row()
    _subhdr(r.cells[0], "Aspecto a Evaluar / Finalidad", C1)
    _subhdr(r.cells[1], "%", C2)
    _subhdr(r.cells[2], "Instrumento de Evaluación", C3)
    _subhdr(r.cells[3], "Momento de Aplicación", C4)

    rows_data = [
        ("1. Evaluación Diagnóstica",
         "Identificar el nivel de conocimientos previos de los participantes como punto de partida del curso.",
         pct_d, _inst(ev.get('instDiagnostica'), 'Cuestionario diagnóstico'), "Al inicio\n(solo referencial)"),
        ("2. Evaluación Formativa",
         "Identificar la comprensión y avance logrado por los participantes durante el curso.",
         pct_f, _inst(ev.get('instFormativa'), 'Lista de cotejo / Guía de observación'), "Intermedia"),
        ("3. Evaluación Final (Sumativa)",
         "Acreditar los aprendizajes adquiridos por los participantes en el proceso de enseñanza-aprendizaje.",
         pct_s, _inst(ev.get('instSumativa'), 'Cuestionario final'), "Al final del curso"),
    ]
    for tipo, finalidad, pct, instrumento, momento in rows_data:
        r = t.add_row()
        _borders(r.cells[0]); _pad(r.cells[0]); _cw(r.cells[0], C1); _clear(r.cells[0])
        _para(r.cells[0], tipo, bold=True, sp_before=20, sp_after=8)
        _para(r.cells[0], f"Finalidad: {finalidad}", size_pt=9, sp_before=8, sp_after=20)
        _val(r.cells[1], f"{pct}%", C2); _val(r.cells[2], instrumento, C3); _val(r.cells[3], momento, C4)

    r = t.add_row(); r.cells[0].merge(r.cells[3])
    _borders(r.cells[0]); _pad(r.cells[0]); _shading(r.cells[0], GRIS_CLARO); _clear(r.cells[0])
    _para(r.cells[0], "d) Criterios de evaluación:", bold=True, sp_before=20, sp_after=8)
    for c in [
        "- Conocimientos Teóricos: Comprende los temas del curso, identifica conceptos clave y los relaciona con su práctica laboral.",
        "- Actitud y comportamiento: Participación activa, respeto a sus compañeros y al instructor.",
        "- Evaluaciones aplicadas: Comprensión de los temas explicados y puntaje aprobatorio en las evaluaciones efectuadas.",
    ]:
        _para(r.cells[0], c, size_pt=9, sp_before=6, sp_after=6)


def _act_paras(cell, text):
    lines = str(text or '').split('\n')
    for i, line in enumerate(lines):
        bold = i == 0 and bool(line.strip())
        _para(cell, line, bold=bold, color=AZUL if bold else NEGRO,
              size_pt=9, sp_before=16 if i == 0 else 6, sp_after=6)


def _tabla_seccion(doc, titulo, filas, primera_col="Etapa"):
    CE = round(TW * .16); CA = round(TW * .46)
    CD = round(TW * .07); CT = round(TW * .15); CM = TW - CE - CA - CD - CT

    t = doc.add_table(rows=0, cols=5)
    _table_setup(t, [CE, CA, CD, CT, CM])

    r = t.add_row(); r.cells[0].merge(r.cells[4])
    _hdr(r.cells[0], titulo, TW)

    r = t.add_row()
    _subhdr(r.cells[0], primera_col, CE); _subhdr(r.cells[1], "Actividades", CA)
    _subhdr(r.cells[2], "Duración", CD);  _subhdr(r.cells[3], "Técnicas Grupales / Instruccionales", CT)
    _subhdr(r.cells[4], "Material y Equipo de Apoyo", CM)

    skip = {'a)', 'b)', 'c)', 'd)', 'e)', 'f)'}
    for fila in filas:
        es_suma = 'suma' in str(fila.get('etapa', '')).lower()
        acts = [a for a in (fila.get('actividades') or []) if str(a or '').strip() not in skip and str(a or '').strip()]

        if es_suma:
            r = t.add_row(); r.cells[0].merge(r.cells[4])
            _borders(r.cells[0]); _pad(r.cells[0], 60, 60, 80, 80)
            _shading(r.cells[0], AZUL_CLARO); _clear(r.cells[0])
            for a in acts:
                _para(r.cells[0], a, bold=True, color=AZUL, size_pt=9, sp_before=20, sp_after=20)
        else:
            r = t.add_row()
            _borders(r.cells[0]); _pad(r.cells[0], 60, 60, 80, 80); _cw(r.cells[0], CE)
            _lines(r.cells[0], fila.get('etapa', ''), size_pt=9)
            _borders(r.cells[1]); _pad(r.cells[1], 60, 60, 80, 80); _cw(r.cells[1], CA); _clear(r.cells[1])
            for a in acts: _act_paras(r.cells[1], a)
            _val_s(r.cells[2], fila.get('duracion', ''), CD)
            _val_s(r.cells[3], fila.get('tecnica', ''), CT)
            _val_s(r.cells[4], fila.get('material', ''), CM)


def _tabla_desarrollo(doc, titulo, filas):
    _tabla_seccion(doc, titulo, filas, primera_col="Temas / Subtemas")
