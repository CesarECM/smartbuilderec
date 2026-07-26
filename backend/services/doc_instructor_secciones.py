import re

from docx import Document
from docx.shared import Pt, RGBColor, Twips

from services.doc_planeacion_xml import AZUL
from services.doc_planeacion_tablas import (
    _tabla_info_general, _tabla_objetivos, _tabla_requerimientos,
    _tabla_evaluaciones, _tabla_seccion, _tabla_desarrollo,
)
from services.doc_planeacion_filas import (
    _filas_previo, _filas_apertura, _filas_desarrollo, _filas_cierre,
)


def _tabla_clave_opcion_multiple(doc: Document, clave_texto: str) -> None:
    pares = re.findall(r'(\d+)\.\s*([A-Da-d])', clave_texto)
    if not pares:
        doc.add_paragraph(clave_texto or "—")
        return
    t = doc.add_table(rows=len(pares) + 1, cols=3)
    t.style = "Table Grid"
    for j, enc in enumerate(["Reactivo", "Respuesta correcta", "Valor"]):
        t.rows[0].cells[j].paragraphs[0].add_run(enc).bold = True
    for i, (num, resp) in enumerate(pares, 1):
        t.rows[i].cells[0].text = num
        t.rows[i].cells[1].text = resp.upper()
        t.rows[i].cells[2].text = "20%"


def _seccion_carta_descriptiva(doc: Document, d: dict) -> None:
    """Embebe la carta descriptiva completa reutilizando los helpers de planeación."""
    doc.add_page_break()

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Twips(0)
    p.paragraph_format.space_after  = Twips(240)
    r = p.add_run("CARTA DESCRIPTIVA DEL CURSO")
    r.font.name = "Arial"; r.font.size = Pt(14); r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(AZUL)

    _tabla_info_general(doc, d)
    doc.add_paragraph().paragraph_format.space_before = Twips(200)
    _tabla_objetivos(doc, d)
    doc.add_paragraph().paragraph_format.space_before = Twips(200)
    _tabla_requerimientos(doc, d)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Twips(80)
    p.paragraph_format.space_after  = Twips(80)
    r = p.add_run(
        "Formas, momentos y criterios de evaluación: La evaluación se llevará a cabo "
        "durante la Apertura, el Desarrollo y el Cierre del Curso/sesión."
    )
    r.font.name = "Arial"; r.font.size = Pt(11); r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(AZUL)
    _tabla_evaluaciones(doc, d)

    doc.add_page_break()
    _tabla_seccion(doc, "PREVIO AL INICIO DEL CURSO — Comprobación de Recursos", _filas_previo(d))
    doc.add_paragraph().paragraph_format.space_before = Twips(200)
    _tabla_seccion(doc, "INICIO DEL CURSO — APERTURA O ENCUADRE", _filas_apertura(d))
    doc.add_page_break()
    _tabla_desarrollo(doc, "DESARROLLO", _filas_desarrollo(d))
    doc.add_page_break()
    _tabla_seccion(doc, "CIERRE", _filas_cierre(d))
    doc.add_page_break()


def _seccion_requerimientos(doc: Document, data) -> None:
    """Sección de requerimientos del lugar usando data.materiales del wizard."""
    mat = data.materiales or {}

    doc.add_heading("REQUERIMIENTOS DEL LUGAR DE CAPACITACIÓN", level=1)

    secciones = [
        (
            "Características del lugar de capacitación",
            mat.get("instalaciones") or
            "Aula iluminada y ventilada, con mesas y sillas suficientes para el número "
            "de participantes, acceso a servicios básicos y salidas de emergencia señalizadas.",
        ),
        (
            "Equipo necesario para desarrollar el curso",
            mat.get("equipo") or
            "Laptop, cañón proyector, extensión eléctrica, pantalla o pared blanca y bocinas.",
        ),
        (
            "Material de apoyo y materiales didácticos",
            mat.get("materialesDidacticos") or
            "Manual del participante, hojas de trabajo, bolígrafos y materiales de apoyo.",
        ),
        (
            "Requerimientos humanos",
            mat.get("humanos") or
            "Instructor certificado y participantes inscritos con el perfil requerido.",
        ),
        (
            "Otros requerimientos",
            mat.get("otros") or "",
        ),
    ]

    for titulo, contenido in secciones:
        if not contenido:
            continue
        doc.add_heading(titulo, level=2)
        for linea in contenido.split("\n"):
            limpia = linea.strip().lstrip("•-–*").strip()
            if limpia:
                doc.add_paragraph(limpia, style="List Bullet")
        doc.add_paragraph()

    doc.add_heading("Recomendaciones de uso del material de apoyo", level=2)
    doc.add_paragraph(
        "El instructor verificará que todos los materiales estén disponibles y en buen estado "
        "antes del inicio del curso. Se recomienda realizar una prueba del equipo audiovisual "
        "con al menos 30 minutos de anticipación para garantizar el desarrollo fluido de la sesión."
    )
    doc.add_paragraph()


def _seccion_instrumentos(doc: Document, ev) -> None:
    """Instrumentos de evaluación completos: instrucciones, reactivos y clave."""
    def _get(attr, *fallbacks):
        for f in (attr, *fallbacks):
            val = (getattr(ev, f, "") or "").strip()
            if val:
                return val
        return ""

    doc.add_page_break()
    doc.add_heading("INSTRUMENTOS DE EVALUACIÓN", level=1)

    # ── Diagnóstica ────────────────────────────────────────────────────────────
    doc.add_heading("EVALUACIÓN DIAGNÓSTICA", level=2)
    header_d = _get("instDiagnosticaHeader")
    if header_d:
        p = doc.add_paragraph()
        p.add_run("Instrucciones para el participante: ").bold = True
        p.add_run(header_d)
    reactivos_d = _get("apfDiagnostica", "instDiagnostica")
    if reactivos_d:
        doc.add_heading("Reactivos", level=3)
        for linea in reactivos_d.split("\n"):
            if linea.strip():
                doc.add_paragraph(linea.strip())
    doc.add_heading("Clave de respuestas", level=3)
    _tabla_clave_opcion_multiple(doc, _get("instDiagnosticaClave"))
    doc.add_paragraph()

    # ── Formativa ──────────────────────────────────────────────────────────────
    doc.add_heading("EVALUACIÓN FORMATIVA", level=2)
    header_f = _get("instFormativaHeader")
    if header_f:
        p = doc.add_paragraph()
        p.add_run("Instrucciones para el participante: ").bold = True
        p.add_run(header_f)
    reactivos_f = _get("apfFormativa", "instFormativa")
    if reactivos_f:
        doc.add_heading("Reactivos / Criterios", level=3)
        for linea in reactivos_f.split("\n"):
            if linea.strip():
                doc.add_paragraph(linea.strip())
    nota_f = _get("notaFormativa")
    if nota_f:
        p = doc.add_paragraph()
        p.add_run("Distribución de puntaje: ").bold = True
        p.add_run(nota_f)
    doc.add_heading("Clave / Criterios de aprobación", level=3)
    doc.add_paragraph(_get("instFormativaClave") or "Ver criterios en el instrumento formativo.")
    doc.add_paragraph()

    # ── Sumativa ───────────────────────────────────────────────────────────────
    doc.add_heading("EVALUACIÓN SUMATIVA", level=2)
    header_s = _get("instSumativaHeader")
    if header_s:
        p = doc.add_paragraph()
        p.add_run("Instrucciones para el participante: ").bold = True
        p.add_run(header_s)
    reactivos_s = _get("apfSumativa", "instSumativa")
    if reactivos_s:
        doc.add_heading("Reactivos", level=3)
        for linea in reactivos_s.split("\n"):
            if linea.strip():
                doc.add_paragraph(linea.strip())
    doc.add_heading("Clave de respuestas", level=3)
    _tabla_clave_opcion_multiple(doc, _get("instSumativaClave"))
    doc.add_paragraph()


def _seccion_unidad_instructor(doc: Document, num: int, temas: list,
                                obj_particular: str, sugerencias: list) -> None:
    """Sección por unidad con tabla de sugerencias didácticas por tema."""
    doc.add_heading(f"UNIDAD {num} — SUGERENCIAS PARA EL DESARROLLO DE TEMAS", level=1)
    doc.add_heading("OBJETIVO PARTICULAR", level=2)
    doc.add_paragraph(obj_particular or f"[Objetivo particular de la unidad {num}]")
    doc.add_paragraph()

    sug_map = {s.get("tema", ""): s for s in (sugerencias or [])}

    for tema in (temas or []):
        doc.add_heading(tema, level=2)
        sug = sug_map.get(tema, {})

        t = doc.add_table(rows=2, cols=3)
        t.style = "Table Grid"
        encabezados = ["Apoyos necesarios", "Actividad / Técnica / Ejemplo", "Evaluación del tema"]
        for j, enc in enumerate(encabezados):
            t.rows[0].cells[j].paragraphs[0].add_run(enc).bold = True

        t.rows[1].cells[0].text = (
            sug.get("apoyos") or "[Describir materiales y recursos de apoyo para este tema]"
        )
        t.rows[1].cells[1].text = (
            sug.get("actividad") or "[Describir la técnica o actividad a aplicar con el grupo]"
        )
        t.rows[1].cells[2].text = (
            sug.get("evaluacion") or "[Indicar criterio, forma y tiempo de evaluación del tema]"
        )
        doc.add_paragraph()

    doc.add_paragraph()
