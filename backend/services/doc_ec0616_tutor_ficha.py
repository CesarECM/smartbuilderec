# ─── services/doc_ec0616_tutor_ficha.py — Ficha de Registro EC0616 Modo Tutor ──
# Clona estructura y formato oficial OC0046: tabla header + tabla 12×8.

import io
import os
import base64

from docx import Document
from docx.shared import Pt, Inches, Cm  # Cm usado en _set_tbl_grid
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

_F_GRIS_OSC = "BFBFBF"   # labels: foto, campos, domicilio, contacto
_F_GRIS_CLR = "F2F2F2"   # filas RENAP: intro, consent, nota pie

_EC = (
    "EC0616: PRESTACIÓN DE SERVICIOS AUXILIARES DE ENFERMERÍA EN CUIDADOS BÁSICOS "
    "Y ORIENTACIÓN A PERSONAS EN UNIDADES DE ATENCIÓN MÉDICA."
)

_RENAP_INTRO = (
    "El Consejo Nacional de Normalización y Certificación de Competencias Laborales (CONOCER) "
    "solicita al candidato la autorización para la publicación de los datos personales a fin de "
    "dar cumplimiento a lo dispuesto en el capítulo séptimo de las Reglas Generales y criterios "
    "para la integración del Sistema Nacional de Competencias, referente al \"Registro Nacional "
    "de Personas Con Competencias Certificadas\" (RENAP) por medio del cual las personas con "
    "competencias certificadas, pueden voluntariamente dar a conocer sus datos personales, para "
    "facilitar su localización, en caso de que organizaciones sindicales, empresas, sector "
    "académico, sector social o público, o alguna otra institución pública o privada, requieran "
    "personal con competencias certificadas en determinada función individual;"
)

_RENAP_CONS = (
    " (   ) doy mi consentimiento al CONOCER para que, en términos del artículo 21 de la Ley "
    "Federal de Transparencia y Acceso a la Información Pública Gubernamental, difunda, "
    "distribuya y publique la información contenida en el documento que se inscribe, para los "
    "propósitos del RENAP. Lo anterior, sin perjuicio de que estoy enterado de que en términos "
    "del artículo 22, fracción III de la misma Ley, no es necesario mi consentimiento respecto "
    "de información que se transmita entre sujetos obligados o entre dependencias y entidades, "
    "cuando los datos respectivos se utilicen para el ejercicio de facultades propias de los mismos."
)

_RENAP_NOTA = (
    "Los datos personales recabados serán protegidos y serán incorporados y tratados en el Sistema "
    "de datos personales RENAP con fundamento en las reglas generales y criterios para integración "
    "y operación del Sistema Nacional de Competencias y cuya finalidad es integrar una base de "
    "datos con información sobre las personas que han obtenido uno o más Certificados de "
    "Competencia, con base en Estándares de Competencia inscritos en el Registro Nacional de "
    "Estándares de Competencia, el cual fue registrado en el Listado de sistemas de Datos "
    "Personales ante el Instituto Federal de Acceso a la Información Pública (www.ifai.org.mx) "
    "y podrán ser trasmitidos a sujetos obligados o dependencias y entidades con la finalidad del "
    "uso en facultades propias de las mismas. Además de otras trasmisiones previstas en Ley. La "
    "Unidad Administrativa responsable del Sistema es el Consejo Nacional de Normalización y "
    "Certificación de Competencias Laborales."
)

_T1_W = [5.5, 2.5, 2.0, 2.0, 1.5, 1.5, 1.5, 1.5]  # 8 cols, total ~18 cm

# Template vacío basado en el documento original (estilos + tema + settings)
_TEMPLATE = os.path.join(os.path.dirname(__file__), "ec0616_ficha_template.docx")


# ── Utilidades de formato ─────────────────────────────────────────────────────

def _docx_bytes(doc: Document) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _chk(val: bool) -> str:
    return "(X)" if val else "(   )"


def _foto_stream(b64: str | None):
    if not b64:
        return None
    try:
        data = b64.split(",", 1)[-1] if "," in b64 else b64
        return io.BytesIO(base64.b64decode(data))
    except Exception:
        return None


def _set_tbl_grid(table, widths_cm: list) -> None:
    tbl = table._tbl
    tblGrid = tbl.find(qn("w:tblGrid"))
    if tblGrid is None:
        tblGrid = OxmlElement("w:tblGrid")
        tbl.insert(0, tblGrid)
    for old in tblGrid.findall(qn("w:gridCol")):
        tblGrid.remove(old)
    for w in widths_cm:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(Cm(w).pt * 20)))
        tblGrid.append(gc)


def _shd(cell, fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    s = OxmlElement("w:shd")
    s.set(qn("w:val"),   "clear")
    s.set(qn("w:color"), "auto")
    s.set(qn("w:fill"),  fill)
    tcPr.append(s)


def _vcenter(cell) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:vAlign")):
        tcPr.remove(old)
    v = OxmlElement("w:vAlign")
    v.set(qn("w:val"), "center")
    tcPr.append(v)


def _cell_margins(cell, dxa: int = 57) -> None:
    """Márgenes internos de celda (padding). El original usa 57 dxa en C0."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn("w:tcMar"))
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for side in ["top", "left", "bottom", "right"]:
        for old in tcMar.findall(qn(f"w:{side}")):
            tcMar.remove(old)
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    str(dxa))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)


def _set_border(cell, side: str, val: str) -> None:
    """Establece un borde de celda. val='nil' lo suprime."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcB = tcPr.find(qn("w:tcBorders"))
    if tcB is None:
        tcB = OxmlElement("w:tcBorders")
        tcPr.append(tcB)
    for old in tcB.findall(qn(f"w:{side}")):
        tcB.remove(old)
    el = OxmlElement(f"w:{side}")
    el.set(qn("w:val"), val)
    if val != "nil":
        el.set(qn("w:color"), "auto")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
    tcB.append(el)


def _nil(cell, *sides: str) -> None:
    """Suprime bordes específicos de una celda."""
    for side in sides:
        _set_border(cell, side, "nil")


def _run(para, text: str, bold: bool = False, size: int = 9,
         italic: bool = False, font: str | None = None) -> None:
    """Agrega un run al párrafo. Solo escribe bold/italic en XML si son True."""
    r = para.add_run(text)
    r.font.size = Pt(size)
    if bold:   r.bold   = True
    if italic: r.italic = True
    if font:
        r.font.name = font
        r._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:cs"), font)


def _w(cell, text: str, bold: bool = False, size: int = 9,
       align=None, italic: bool = False, font: str | None = None) -> None:
    """Escribe en el párrafo inicial de la celda (sin forzar spacing)."""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    if align is not None:
        p.alignment = align
    _run(p, text, bold=bold, size=size, italic=italic, font=font)


def _wa(cell, text: str, bold: bool = False, size: int = 9,
        align=None, italic: bool = False, font: str | None = None) -> None:
    """Agrega un párrafo adicional a la celda."""
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    if align is not None:
        p.alignment = align
    if text:
        _run(p, text, bold=bold, size=size, italic=italic, font=font)


def _line_spacing(cell, line: int = 240, rule: str = "auto") -> None:
    """Establece interlineado en todos los párrafos de la celda."""
    for p in cell.paragraphs:
        pPr = p._p.get_or_add_pPr()
        spc = pPr.find(qn("w:spacing"))
        if spc is None:
            spc = OxmlElement("w:spacing")
            pPr.append(spc)
        spc.set(qn("w:line"),     str(line))
        spc.set(qn("w:lineRule"), rule)


# ── Sección Información Confidencial ─────────────────────────────────────────

def _gen_confidencial(doc: Document, ficha: dict, conf: dict) -> None:
    h = doc.add_heading("", level=2)
    h.add_run("Información Confidencial:").font.size = Pt(11)

    def _row(text: str, bold: bool = False):
        p = doc.add_paragraph(text)
        p.runs[0].font.size = Pt(9)
        p.runs[0].bold = bold
        p.paragraph_format.space_after = Pt(3)

    p_inst = doc.add_paragraph('Marca con una "x" en el recuadro de la respuesta elegida.')
    p_inst.runs[0].font.size = Pt(9)

    lee      = conf.get("leer_escribir")   == "si"
    estudios = conf.get("tiene_estudios")  == "si"
    _row(
        f"¿Sabe Leer y Escribir? {_chk(lee)} Sí {_chk(not lee)} No     "
        f"¿Cuenta con Estudios? {_chk(estudios)} Sí {_chk(not estudios)} No     "
        f"Cuáles: {conf.get('estudios_cuales') or '—'}     "
        f"Último grado: {conf.get('ultimo_grado') or '—'}"
    )
    disc  = conf.get("tiene_discapacidad") == "si"
    tipos = conf.get("discapacidad_tipo") or []
    _row(f"¿Tiene algún tipo de Discapacidad? {_chk(disc)} Sí {_chk(not disc)} No")
    _row("Cual: " + "  ".join(
        f"{_chk(t in tipos)} {t.capitalize()}"
        for t in ["motriz", "visual", "auditiva", "lenguaje", "intelectual", "otras"]
    ))
    _row(f"¿Qué Idioma(s) o lengua(s) habla? {conf.get('idiomas') or '—'}")
    trabaja = conf.get("trabaja_actualmente") == "si"
    _row(
        f"¿Trabaja Actualmente? {_chk(trabaja)} Sí {_chk(not trabaja)} No     "
        f"Puesto de Trabajo: {conf.get('puesto_trabajo') or '—'}"
    )
    if trabaja:
        _row(
            f"Empresa/Institución: {conf.get('empresa_trabajo') or '—'}     "
            f"Tel: {conf.get('telefono_trabajo') or '—'}"
        )
    _row("Experiencia Laboral:", bold=True)
    for n in range(1, 4):
        _row(
            f"  {n}. {conf.get(f'exp{n}_empresa') or '—'} | "
            f"{conf.get(f'exp{n}_periodo') or '—'} | "
            f"{conf.get(f'exp{n}_actividades') or '—'}"
        )
    _row(f"Observaciones: {conf.get('observaciones') or '—'}")
    cert = conf.get("tiene_certificacion") == "si"
    _row(
        f"¿Cuenta con alguna Certificación? {_chk(cert)} Sí {_chk(not cert)} No     "
        f"Cuáles: {conf.get('certificacion_cuales') or '—'}"
    )
    doc.add_paragraph()
    p_decl = doc.add_paragraph(
        "DECLARO BAJO PROTESTA DE DECIR VERDAD QUE LOS DATOS ASENTADOS EN ESTE "
        "DOCUMENTO SON CORRECTOS Y VERDADEROS."
    )
    p_decl.runs[0].bold = True
    p_decl.runs[0].font.size = Pt(9)
    doc.add_paragraph()
    doc.add_paragraph("Atentamente").runs[0].font.size = Pt(9)
    doc.add_paragraph()
    nombre_c = f"{ficha.get('nombre_candidato', '')} {ficha.get('apellidos_candidato', '')}".strip()
    doc.add_paragraph(f"({nombre_c or 'Nombre y firma del Candidato'})").runs[0].font.size = Pt(9)
    doc.add_paragraph()
    nota_p = doc.add_paragraph(
        "Nota: Esta información se debe mantener en el portafolio de evidencias del candidato"
    )
    nota_p.runs[0].font.size = Pt(8)
    nota_p.runs[0].italic = True


# ── Función principal ─────────────────────────────────────────────────────────

def generar_ficha(ficha: dict, conf: dict) -> bytes:
    # Hereda estilos, tema, settings y márgenes del original
    doc = Document(_TEMPLATE)

    # ── Tabla 0: Encabezado ───────────────────────────────────────────────────
    t0 = doc.add_table(rows=1, cols=5)
    t0.style = "Table Grid"
    _set_tbl_grid(t0, [3.3, 8.5, 0.5, 2.5, 3.2])
    c = t0.rows[0].cells
    _shd(c[0], _F_GRIS_OSC)
    _w(c[0], "Estándar de Competencia:", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _w(c[1], _EC, size=10)
    _shd(c[3], _F_GRIS_OSC)
    _w(c[3], "Fecha:", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _w(c[4], ficha.get("fecha_evaluacion") or "—")

    # ── Párrafo "Datos Personales:" ───────────────────────────────────────────
    p_dp = doc.add_paragraph()
    r_dp = p_dp.add_run("Datos Personales:")
    r_dp.bold = True
    r_dp.font.size = Pt(10)
    p_dp.paragraph_format.space_before = Pt(4)
    p_dp.paragraph_format.space_after  = Pt(2)

    # ── Tabla 1: Principal (12 filas × 8 cols) ────────────────────────────────
    t1 = doc.add_table(rows=12, cols=8)
    t1.style = "Table Grid"
    _set_tbl_grid(t1, _T1_W)

    # ── R0: Intro CONOCER ─────────────────────────────────────────────────────
    t1.cell(0, 0).merge(t1.cell(0, 7))
    c0 = t1.cell(0, 0)
    _shd(c0, _F_GRIS_CLR)
    _vcenter(c0)
    _nil(c0, "bottom")                              # borde inferior suprimido
    _w(c0, _RENAP_INTRO, size=9, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # ── R1: separador vacío ───────────────────────────────────────────────────
    t1.cell(1, 1).merge(t1.cell(1, 7))

    # ── R2-R10 C0: Consentimiento RENAP (fusión vertical) ────────────────────
    t1.cell(2, 0).merge(t1.cell(10, 0))
    cc = t1.cell(2, 0)
    _shd(cc, _F_GRIS_CLR)
    _cell_margins(cc, 57)                           # 57 dxa = margen original
    _nil(cc, "top")                                 # borde superior suprimido

    # p0: SI/NO en runs separados (bold/normal), 8pt, JUSTIFY
    cons = ficha.get("consentimiento_renap", "no") == "si"
    p0 = cc.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p0.paragraph_format.space_after = Pt(0)
    _run(p0, "SI ",        bold=True,  size=8)
    _run(p0, _chk(cons),               size=8)
    _run(p0, "  NO",       bold=True,  size=8)
    _run(p0, _RENAP_CONS,              size=8)

    # p1, p2: párrafos vacíos de separación (como en el original)
    _wa(cc, "")
    _wa(cc, "")

    # p3, p4: firma
    _wa(cc, "_____________________________", size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    _wa(cc, "Nombre y Firma",               size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── R2-R6 C1: Fotografía (fusión vertical) ────────────────────────────────
    t1.cell(2, 1).merge(t1.cell(6, 1))
    fc = t1.cell(2, 1)
    _shd(fc, _F_GRIS_OSC)
    _vcenter(fc)
    p_lbl = fc.paragraphs[0]
    p_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_lbl.paragraph_format.space_after = Pt(4)
    _run(p_lbl, "Fotografía Digital ", bold=True,  size=10)
    _run(p_lbl, "(Reciente)",           bold=False, size=10)
    foto = _foto_stream(ficha.get("foto_base64"))
    if foto:
        p_img = fc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_after = Pt(0)
        p_img.add_run().add_picture(foto, width=Inches(0.9), height=Inches(1.2))

    # ── Helper: celda label gris ──────────────────────────────────────────────
    def _label(row: int, text: str, align=WD_ALIGN_PARAGRAPH.RIGHT) -> None:
        t1.cell(row, 2).merge(t1.cell(row, 3))
        lc = t1.cell(row, 2)
        _shd(lc, _F_GRIS_OSC)
        _vcenter(lc)
        _w(lc, text, bold=True, size=10, align=align)

    def _value(row: int):
        t1.cell(row, 4).merge(t1.cell(row, 7))
        vc = t1.cell(row, 4)
        _vcenter(vc)
        return vc

    # ── R2-R6: Datos personales ───────────────────────────────────────────────
    _label(2, "Nombre Completo:")
    nombre = f"{ficha.get('nombre_candidato', '')} {ficha.get('apellidos_candidato', '')}".strip()
    _w(_value(2), nombre or "—", size=10)

    _label(3, "Lugar de Nacimiento:")
    _w(_value(3), ficha.get("lugar_nacimiento") or "—", size=10)

    _label(4, "Nacionalidad:")
    _w(_value(4), ficha.get("nacionalidad") or "Mexicana", size=10)

    _label(5, "CURP:")
    _w(_value(5), ficha.get("curp") or "—", size=10)

    # R6: Género + Fecha Nacimiento
    _label(6, "Género:")
    genero = ficha.get("genero", "")
    gc6 = t1.cell(6, 4)
    _vcenter(gc6)
    _w(gc6, f"{_chk(genero == 'H')} H   {_chk(genero == 'M')} M", size=10)

    t1.cell(6, 5).merge(t1.cell(6, 6))
    fn6 = t1.cell(6, 5)
    _shd(fn6, _F_GRIS_OSC)
    _vcenter(fn6)
    _w(fn6, "Fecha de Nacimiento:", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    vc7 = t1.cell(6, 7)
    _vcenter(vc7)
    _w(vc7, ficha.get("fecha_nacimiento") or "—", size=10)

    # ── R7: Header Domicilio ──────────────────────────────────────────────────
    t1.cell(7, 1).merge(t1.cell(7, 7))
    dc = t1.cell(7, 1)
    _shd(dc, _F_GRIS_OSC)
    _vcenter(dc)
    _w(dc, "Domicilio Particular", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── R8: Calle / Núm / CP / Colonia ───────────────────────────────────────
    t1.cell(8, 1).merge(t1.cell(8, 7))
    r8c = t1.cell(8, 1)
    _vcenter(r8c)
    _w(r8c,
       f"Calle: {ficha.get('domicilio_calle') or '—'}   "
       f"Núm: {ficha.get('domicilio_numero') or '—'}   "
       f"C.P.: {ficha.get('domicilio_cp') or '—'}   "
       f"Colonia: {ficha.get('domicilio_colonia') or '—'}",
       size=10)

    # ── R9: Municipio / Entidad ───────────────────────────────────────────────
    t1.cell(9, 1).merge(t1.cell(9, 2))
    t1.cell(9, 3).merge(t1.cell(9, 5))
    t1.cell(9, 6).merge(t1.cell(9, 7))
    _vcenter(t1.cell(9, 1))
    _vcenter(t1.cell(9, 3))
    _w(t1.cell(9, 1), f"Municipio: {ficha.get('domicilio_ciudad') or '—'}",           size=10)
    _w(t1.cell(9, 3), f"Entidad Federativa: {ficha.get('domicilio_entidad') or '—'}", size=10)

    # ── R10: E-mail / Teléfono / Celular ─────────────────────────────────────
    t1.cell(10, 1).merge(t1.cell(10, 2))
    t1.cell(10, 3).merge(t1.cell(10, 5))
    t1.cell(10, 6).merge(t1.cell(10, 7))

    # Label + valor en la misma celda gris (label bold, valor normal)
    ec = t1.cell(10, 1)
    _shd(ec, _F_GRIS_OSC)
    _vcenter(ec)
    _nil(ec, "right")                               # sin borde derecho
    _w(ec,  "E-mail",                  bold=True,  size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _wa(ec, ficha.get("email") or "—",             size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    tc10 = t1.cell(10, 3)
    _shd(tc10, _F_GRIS_OSC)
    _vcenter(tc10)
    _nil(tc10, "left", "right")                     # sin bordes laterales
    _w(tc10,  "Teléfono",                 bold=True,  size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _wa(tc10, ficha.get("telefono") or "—",          size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    cc10 = t1.cell(10, 6)
    _shd(cc10, _F_GRIS_OSC)
    _vcenter(cc10)
    _nil(cc10, "left")                              # sin borde izquierdo
    _w(cc10,  "Teléfono Celular",  bold=True,  size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _wa(cc10, "—",                             size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── R11: Nota RENAP pie ───────────────────────────────────────────────────
    t1.cell(11, 0).merge(t1.cell(11, 7))
    nc = t1.cell(11, 0)
    _shd(nc, _F_GRIS_CLR)
    _w(nc, _RENAP_NOTA, size=8, font="Arial")
    _line_spacing(nc, line=240, rule="auto")        # interlineado simple original

    # ── Información Confidencial (página 2) ───────────────────────────────────
    doc.add_page_break()
    _gen_confidencial(doc, ficha, conf)

    return _docx_bytes(doc)
