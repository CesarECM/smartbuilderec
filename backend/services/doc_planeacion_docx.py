import io
import os
import tempfile

from docx import Document
from docx.shared import Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH

from services.doc_planeacion_xml import AZUL, TW, _run
from services.doc_planeacion_tablas import (
    _tabla_info_general, _tabla_objetivos, _tabla_requerimientos,
    _tabla_evaluaciones,
)
from services.doc_planeacion_filas import (
    _filas_previo, _filas_apertura, _filas_desarrollo, _filas_cierre,
)
from services.doc_planeacion_tablas import _tabla_seccion, _tabla_desarrollo


def _preprocesar(payload: dict) -> dict:
    """Normaliza campos legacy (rhObjetivo/rhInstrucciones → rhDetalle) para retrocompatibilidad."""
    tc = dict(payload.get('tecnicas') or {})

    rh_objetivo      = tc.get('rhObjetivo', '')
    rh_instrucciones = tc.get('rhInstrucciones', '')
    en_objetivo      = tc.get('enObjetivo', '')
    en_instrucciones = tc.get('enInstrucciones', '')

    if not tc.get('rhDetalle') and (rh_objetivo or rh_instrucciones):
        tc['rhDetalle'] = (
            f"a) Explicará objetivo de la técnica:\n{rh_objetivo}\n\n"
            f"b) Dará las instrucciones de la técnica:\n{rh_instrucciones}\n\n"
            f"c) Mencionará el tiempo para realizarla.\n\n"
            f"d) Propiciará la participación del grupo.\n\n"
            f"e) Integrará al grupo.\n\n"
            f"f) Controlará el tiempo."
        )
    if not tc.get('enDetalle') and (en_objetivo or en_instrucciones):
        tc['enDetalle'] = (
            f"a) Explicará objetivo de la técnica:\n{en_objetivo}\n\n"
            f"b) Dará las instrucciones de la técnica:\n{en_instrucciones}"
        )

    return {**payload, 'tecnicas': tc}


def generar_planeacion_docx(payload: dict, branding: dict | None = None) -> bytes:
    d   = _preprocesar(payload)
    br  = branding or {}
    empresa_footer = br.get("empresa") or "SmartBuilder EC  •  Centro ECM"
    logo_url       = br.get("logo_url")

    doc = Document()
    sec = doc.sections[0]
    sec.page_width  = Twips(12240); sec.page_height = Twips(15840)
    sec.left_margin = sec.right_margin = Twips(720)
    sec.top_margin  = sec.bottom_margin = Twips(720)
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(10)

    if logo_url:
        try:
            import urllib.request
            with urllib.request.urlopen(logo_url, timeout=5) as resp:
                logo_bytes = resp.read()
            ext = logo_url.split(".")[-1].lower().split("?")[0]
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}")
            tmp.write(logo_bytes); tmp.close()
            lp = doc.add_paragraph()
            lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            lp.paragraph_format.space_before = Twips(0)
            lp.paragraph_format.space_after  = Twips(80)
            lp.add_run().add_picture(tmp.name, height=Pt(36))
            os.unlink(tmp.name)
        except Exception as e_logo:
            print(f"[branding] No se pudo insertar logo: {e_logo}")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Twips(0)
    p.paragraph_format.space_after  = Twips(240)
    r = p.add_run("DOCUMENTO DE PLANEACIÓN DEL CURSO / CARTA DESCRIPTIVA")
    r.font.name = 'Arial'; r.font.size = Pt(14); r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(AZUL)

    _tabla_info_general(doc, d)
    doc.add_paragraph().paragraph_format.space_before = Twips(200)
    _tabla_objetivos(doc, d)
    doc.add_paragraph().paragraph_format.space_before = Twips(200)
    _tabla_requerimientos(doc, d)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Twips(80)
    p.paragraph_format.space_after  = Twips(80)
    r = p.add_run("Formas, momentos y criterios de evaluación: La evaluación se llevará a cabo durante la Apertura, el Desarrollo y el Cierre del Curso/sesión.")
    r.font.name = 'Arial'; r.font.size = Pt(11); r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(AZUL)
    _tabla_evaluaciones(doc, d)

    doc.add_page_break()

    _tabla_seccion(doc, "PREVIO AL INICIO DEL CURSO — Comprobación de Recursos", _filas_previo(d))
    doc.add_paragraph().paragraph_format.space_before = Twips(200)
    _tabla_seccion(doc, "INICIO DEL CURSO — APERTURA O ENCUADRE", _filas_apertura(d))
    doc.add_page_break()
    _tabla_desarrollo(doc, "DESARROLLO", _filas_desarrollo(d))
    doc.add_page_break()
    _tabla_seccion(doc, "CIERRE", _filas_cierre(d))

    dat = d.get('datos') or {}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Twips(320)
    r = p.add_run(f"{empresa_footer}  •  {dat.get('nombreCurso','')}  •  {dat.get('fecha','')}")
    r.font.name = 'Arial'; r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string("999999")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
