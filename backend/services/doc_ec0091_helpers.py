import io
import zipfile
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

_AZUL = RGBColor(0x1A, 0x4A, 0x6B)


def _titulo(doc: Document, texto: str) -> None:
    h = doc.add_heading("", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run(texto)
    r.font.color.rgb = _AZUL
    r.font.size = Pt(14)


def _encabezado_ec0091(doc: Document, datos: dict) -> None:
    doc.add_paragraph(f"Organismo Certificador: {datos.get('oc91NombreOrganismo', '')} — Clave: {datos.get('oc91Clave', '')}")
    doc.add_paragraph(f"Verificador Externo: {datos.get('ve91Nombre', '')} — No. Cert.: {datos.get('ve91NoCert', '')}")
    tipo = datos.get('ent91Tipo', 'CE')
    doc.add_paragraph(f"Entidad verificada ({tipo}): {datos.get('ent91Nombre', '')} — Clave: {datos.get('ent91Clave', '')}")
    doc.add_paragraph(f"Responsable que recibe: {datos.get('resp91Nombre', '')} — Cargo: {datos.get('resp91Cargo', '')}")
    doc.add_paragraph(f"Fecha de verificación: {datos.get('fecha91Verificacion', '')}")
    doc.add_paragraph()


def _docx_bytes(doc: Document) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def generar_plan_verificacion(payload: dict) -> bytes:
    doc = Document()
    _titulo(doc, "Plan de Verificación Externa — EC0091")
    _titulo(doc, "Documento 01")
    doc.add_paragraph()
    _encabezado_ec0091(doc, payload.get("datos", {}))

    obj = payload.get("objetivo", {})
    doc.add_heading("Objetivo de la Verificación", level=2)
    doc.add_paragraph(obj.get("objetivo", ""))
    doc.add_heading("Alcance", level=2)
    doc.add_paragraph(obj.get("alcance", ""))

    ant = payload.get("antecedentes", {})
    doc.add_heading("Antecedentes", level=2)
    tiene = ant.get("ant91TieneAntecedentes", "no")
    doc.add_paragraph(f"¿Verificaciones previas?: {'Sí' if tiene == 'si' else 'No'}")
    if tiene == "si":
        doc.add_paragraph(f"Verificaciones previas: {ant.get('ant91Previas', '')}")
        doc.add_paragraph(f"Acciones correctivas: {ant.get('ant91AccCorrectivas', '')}")
    if ant.get("ant91Observaciones"):
        doc.add_paragraph(f"Observaciones: {ant.get('ant91Observaciones', '')}")

    lin = payload.get("lineas", {}).get("lineas", {})
    doc.add_heading("Líneas de Verificación Seleccionadas", level=2)
    for tipo, datos in lin.items():
        if datos.get("seleccionada"):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{tipo.capitalize()}: ").bold = True
            p.add_run(datos.get("actividades", ""))

    mues = payload.get("muestreo", {})
    doc.add_heading("Muestreo Aleatorio Simple", level=2)
    doc.add_paragraph(f"Total instrumentos (N): {mues.get('N', '')}")
    doc.add_paragraph(f"Tamaño de muestra: {mues.get('muestra', '')}")
    doc.add_paragraph(f"Número de aceptación: {mues.get('aceptacion', '')}")
    doc.add_paragraph(f"Número de rechazo: {mues.get('rechazo', '')}")
    if mues.get("explicacion"):
        doc.add_paragraph(f"Interpretación: {mues.get('explicacion', '')}")

    cron = payload.get("cronograma", {}).get("actividades", [])
    if cron:
        doc.add_heading("Cronograma de Actividades", level=2)
        tabla = doc.add_table(rows=1, cols=5)
        tabla.style = "Table Grid"
        for i, col in enumerate(["Fecha", "Hora", "Actividad", "Responsable", "Lugar"]):
            tabla.rows[0].cells[i].text = col
        for act in cron:
            row = tabla.add_row()
            for i, campo in enumerate(["fecha", "hora", "actividad", "responsable", "lugar"]):
                row.cells[i].text = act.get(campo, "")

    return _docx_bytes(doc)


def generar_lista_verificacion(payload: dict) -> bytes:
    doc = Document()
    _titulo(doc, "Lista de Verificación — EC0091")
    _titulo(doc, "Documento 02")
    doc.add_paragraph()
    _encabezado_ec0091(doc, payload.get("datos", {}))
    doc.add_heading("Criterios y Preguntas de Verificación", level=2)
    lista = payload.get("lista", {}).get("preguntas", "")
    for linea in lista.split("\n"):
        doc.add_paragraph(linea)
    return _docx_bytes(doc)


def generar_lista_aplicada(payload: dict) -> bytes:
    doc = Document()
    _titulo(doc, "Lista de Verificación Aplicada — EC0091")
    _titulo(doc, "Documento 03 — Registro de Campo")
    doc.add_paragraph()
    _encabezado_ec0091(doc, payload.get("datos", {}))
    ejec = payload.get("ejecucion", {})
    doc.add_heading("Registro de Cumplimientos", level=2)
    for linea in ejec.get("registro", "").split("\n"):
        doc.add_paragraph(linea)
    if ejec.get("clasificacion"):
        doc.add_heading("Clasificación de Incumplimientos", level=2)
        for linea in ejec.get("clasificacion", "").split("\n"):
            doc.add_paragraph(linea)
    return _docx_bytes(doc)


def generar_reporte_hallazgos(payload: dict) -> bytes:
    doc = Document()
    _titulo(doc, "Reporte de Hallazgos y Acciones Correctivas — EC0091")
    _titulo(doc, "Documento 04")
    doc.add_paragraph()
    _encabezado_ec0091(doc, payload.get("datos", {}))
    hallazgos = payload.get("hallazgos", {}).get("hallazgos", [])
    for i, h in enumerate(hallazgos, 1):
        doc.add_heading(f"Hallazgo {i}", level=2)
        doc.add_paragraph(f"Descripción: {h.get('descripcion', '')}")
        doc.add_paragraph(f"Causa raíz: {h.get('causa', '')}")
        doc.add_paragraph(f"Acción correctiva: {h.get('accion', '')}")
        doc.add_paragraph(f"Indicador: {h.get('indicador', '')}")
        doc.add_paragraph(f"Plazo: {h.get('plazo', '')}")
        doc.add_paragraph("Firma responsable: ___________________")
        doc.add_paragraph()
    return _docx_bytes(doc)


def generar_informe_verificacion(payload: dict) -> bytes:
    doc = Document()
    _titulo(doc, "Informe de Verificación Externa — EC0091")
    _titulo(doc, "Documento 05")
    doc.add_paragraph()
    _encabezado_ec0091(doc, payload.get("datos", {}))
    inf = payload.get("informe", {})
    doc.add_heading("Resumen Ejecutivo", level=2)
    doc.add_paragraph(inf.get("inf91Resumen", ""))
    doc.add_heading("Conclusiones", level=2)
    doc.add_paragraph(inf.get("inf91Conclusiones", ""))
    doc.add_heading("Resultado Cuantitativo", level=2)
    doc.add_paragraph(inf.get("inf91ResultadoFinal", ""))
    doc.add_heading("Dictamen", level=2)
    dictamen_labels = {
        "aprobado": "APROBADO — Cumple con los requisitos del EC0091",
        "condicionado": "CONDICIONADO — Requiere implementar acciones correctivas",
        "no_aprobado": "NO APROBADO — Incumplimientos críticos detectados",
        "pendiente": "PENDIENTE DE DICTAMEN",
    }
    doc.add_paragraph(dictamen_labels.get(inf.get("inf91Dictamen", "pendiente"), ""))
    cierre = payload.get("cierre", {})
    doc.add_heading("Firmas", level=2)
    doc.add_paragraph(f"Fecha de entrega del informe: {cierre.get('cierre91FechaEntrega', '')}")
    doc.add_paragraph("Firma del Verificador Externo: ___________________")
    doc.add_paragraph("Firma del Responsable de la Entidad: ___________________")
    doc.add_paragraph(f"Observaciones de cierre: {cierre.get('cierre91Observaciones', '')}")
    return _docx_bytes(doc)


def generar_zip_ec0091(payload: dict) -> bytes:
    docs = [
        ("01_Plan_Verificacion_Externa.docx",     generar_plan_verificacion(payload)),
        ("02_Lista_Verificacion.docx",            generar_lista_verificacion(payload)),
        ("03_Lista_Verificacion_Aplicada.docx",   generar_lista_aplicada(payload)),
        ("04_Reporte_Hallazgos.docx",             generar_reporte_hallazgos(payload)),
        ("05_Informe_Verificacion_Externa.docx",  generar_informe_verificacion(payload)),
    ]
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
        for nombre, contenido in docs:
            zf.writestr(nombre, contenido)
    buf.seek(0)
    return buf.read()
