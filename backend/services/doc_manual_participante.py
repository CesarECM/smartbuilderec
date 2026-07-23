import json
import os

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openai import OpenAI

from services.doc_helpers import load_prompt, _docx_bytes, _titulo

_client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
_MODEL  = os.getenv("OPENAI_MODEL_GRL", "gpt-4o-mini")


def _ia_contenido(data) -> dict:
    datos = data.datos
    obj   = data.objetivos
    tem   = data.temario

    def _lista(items): return ", ".join(i for i in (items or []) if i) or "—"

    prompt = load_prompt(
        "manual_participante_prompt.txt",
        nombreCurso=datos.nombreCurso or "",
        objetivoGeneral=obj.general or "",
        objetivoCognitivo=obj.cognitiva or "",
        objetivoPsicomotriz=obj.psicomotriz or "",
        objetivoAfectivo=obj.afectiva or "",
        beneficios=getattr(data, "beneficios", "") or "",
        temasU1=_lista(tem.u1),
        temasU2=_lista(tem.u2),
        temasU3=_lista(tem.u3),
    )
    resp = _client.chat.completions.create(
        model=_MODEL,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.4,
        response_format={"type": "json_object"},
    )
    return json.loads(resp.choices[0].message.content.strip())


def _heading(doc: Document, texto: str, level: int = 1) -> None:
    h = doc.add_heading(texto, level=level)
    h.runs[0].font.color.rgb = RGBColor(0x1F, 0x3B, 0x6D)


def _tabla_criterios(doc: Document, ev) -> None:
    def _pct(a, b): return a or b or 0
    tabla = doc.add_table(rows=4, cols=2)
    tabla.style = "Table Grid"
    encabezados = ["Instrumentos", "PORCENTAJE"]
    for i, txt in enumerate(encabezados):
        c = tabla.rows[0].cells[i]
        c.text = txt
        c.paragraphs[0].runs[0].bold = True

    filas = [
        ("Evaluación diagnóstica (al inicio)",      f"{_pct(ev.pctDiagnostica, ev.pctDiag)} %"),
        ("Evaluación formativa (intermedio)",        f"{_pct(ev.pctFormativa,   ev.pctForm)} %"),
        ("Evaluación sumativa (al final del curso)", f"{_pct(ev.pctSumativa,    ev.pctSuma)} %"),
    ]
    for i, (inst, pct) in enumerate(filas, 1):
        tabla.rows[i].cells[0].text = inst
        tabla.rows[i].cells[1].text = pct


def _tabla_expectativas(doc: Document) -> None:
    tabla = doc.add_table(rows=5, cols=2)
    tabla.style = "Table Grid"
    tabla.rows[0].cells[0].text = "Expectativas"
    tabla.rows[0].cells[1].text = "Compromisos como participante"
    for fila in tabla.rows:
        for cell in fila.cells:
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].bold = True
            break


def _seccion_unidad(doc: Document, num: int, temas: list, obj_particular: str,
                    conclusion: str) -> None:
    _heading(doc, f"PARTE {num}:", level=1)
    _heading(doc, "OBJETIVO PARTICULAR", level=2)
    doc.add_paragraph(obj_particular or f"[Objetivo particular de la unidad {num}]")
    doc.add_paragraph()

    for i, tema in enumerate(temas, 1):
        _heading(doc, f"{num}.{i} {tema}", level=2)
        doc.add_paragraph("[Desarrollar el contenido de este tema]")
        doc.add_paragraph()

    _heading(doc, f"{num}.{len(temas)+1} CONCLUSIONES", level=2)
    doc.add_paragraph(conclusion or f"[Redactar conclusión de la unidad {num}]")
    doc.add_paragraph()

    _heading(doc, "Actividad de reforzamiento de aprendizaje:", level=2)
    doc.add_paragraph(
        "Se realizará una actividad para verificar la comprensión de los temas "
        "desarrollados en esta parte del curso."
    )
    doc.add_paragraph()


def generar_manual_participante(data) -> bytes:
    ia = _ia_contenido(data)

    datos = data.datos
    obj   = data.objetivos
    tem   = data.temario
    ev    = data.evaluaciones
    cierre = data.cierre or {}

    doc = Document()
    _titulo(doc, "MANUAL DEL PARTICIPANTE", subtitulo=datos.nombreCurso or "")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Nombre del Instructor: {datos.instructor or ''}").bold = True
    doc.add_paragraph()

    _heading(doc, "ÍNDICE")
    doc.add_paragraph("[El índice se genera automáticamente en Word: Referencias → Tabla de contenido]")
    doc.add_paragraph()

    _heading(doc, "PRESENTACIÓN")
    doc.add_paragraph(ia.get("presentacion", ""))
    doc.add_paragraph()

    _heading(doc, "BIENVENIDA")
    doc.add_paragraph(
        f"Bienvenido al curso «{datos.nombreCurso or 'nombre del curso'}». "
        "Durante el desarrollo del programa adquirirás los conocimientos y habilidades "
        "necesarios para alcanzar los objetivos planteados. Esperamos que aproveches "
        "al máximo cada uno de los temas y actividades propuestos."
    )
    doc.add_paragraph()

    _heading(doc, "RECOMENDACIONES DE USO DEL MANUAL")
    doc.add_paragraph(
        "Te recomendamos leer y analizar cada uno de los temas de este manual durante "
        "el desarrollo del curso, con el fin de comprenderlos y sacar el mayor provecho "
        "posible, logrando así los objetivos planteados y aprovechando el tiempo al máximo."
    )
    doc.add_paragraph()

    _heading(doc, "ORGANIZACIÓN DEL MANUAL")
    doc.add_paragraph(
        f"Este manual está integrado por un índice, una presentación, una introducción, "
        f"los objetivos del curso «{datos.nombreCurso or ''}», el desarrollo de los temas "
        f"organizados en tres partes, conclusiones por unidad y un apartado de referencias bibliográficas."
    )
    doc.add_paragraph()

    _heading(doc, "INTRODUCCIÓN")
    doc.add_paragraph(ia.get("introduccion", ""))
    doc.add_paragraph()

    _heading(doc, "BENEFICIOS QUE EL CURSO APORTARÁ A LOS PARTICIPANTES")
    doc.add_paragraph(getattr(data, "beneficios", "") or "[Ver documento de planeación]")
    doc.add_paragraph()

    _heading(doc, "ENFOQUE DIDÁCTICO DEL CURSO")
    doc.add_paragraph(
        "La modalidad de este curso de formación se basa en un esquema presencial/híbrido, "
        "donde el participante aprende participando y realizando las actividades de "
        "aprendizaje propuestas."
    )
    doc.add_paragraph()

    _heading(doc, "OBJETIVO GENERAL DEL CURSO")
    doc.add_paragraph(obj.general or "")
    doc.add_paragraph()

    _heading(doc, "OBJETIVOS PARTICULARES", level=2)
    for o in [obj.cognitiva, obj.psicomotriz, obj.afectiva]:
        if o:
            doc.add_paragraph(o, style="List Bullet")
    doc.add_paragraph()

    _heading(doc, "ANÁLISIS DE EXPECTATIVAS Y ESTABLECIMIENTO DE COMPROMISOS")
    doc.add_paragraph("Instrucciones:")
    doc.add_paragraph("1. Conteste de forma individual cuáles son sus expectativas y compromisos del curso.")
    doc.add_paragraph("2. Comente sus respuestas y anote sus comentarios.")
    doc.add_paragraph()
    _tabla_expectativas(doc)
    doc.add_paragraph()

    _heading(doc, "CRITERIOS DE EVALUACIÓN")
    doc.add_paragraph(
        "La calificación mínima aprobatoria para este curso es del 80%. "
        "Esta calificación será distribuida entre los siguientes factores:"
    )
    doc.add_paragraph()
    _tabla_criterios(doc, ev)
    doc.add_paragraph()

    objetivos_particulares = [obj.cognitiva, obj.psicomotriz, obj.afectiva]
    unidades = [
        (tem.u1, ia.get("conclusion_u1", "")),
        (tem.u2, ia.get("conclusion_u2", "")),
        (tem.u3, ia.get("conclusion_u3", "")),
    ]
    for i, (temas, conclusion) in enumerate(unidades, 1):
        obj_part = objetivos_particulares[i - 1] if i - 1 < len(objetivos_particulares) else ""
        _seccion_unidad(doc, i, temas or [], obj_part, conclusion)

    _heading(doc, "REFERENCIAS BIBLIOGRÁFICAS")
    refs = cierre.get("referencias", "")
    doc.add_paragraph(refs if refs else "[Incluir referencias bibliográficas del curso]")

    return _docx_bytes(doc)
