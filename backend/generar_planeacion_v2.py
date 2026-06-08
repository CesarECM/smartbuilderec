"""
Generador del Documento de Planeación EC0217.01 — Python puro (D#4)
Reemplaza el subproceso Node.js de generar_planeacion.js
"""
from docx import Document
from docx.shared import Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io

AZUL       = "1F3B6D"
AZUL_CLARO = "D6E4F0"
GRIS_CLARO = "F2F2F2"
AMARILLO   = "FFF9E6"
BLANCO     = "FFFFFF"
NEGRO      = "000000"
TW         = 10800  # ancho útil carta en twips

# ─── XML helpers ──────────────────────────────────────────────────────────────

def _shading(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    for e in tcPr.findall(qn('w:shd')): tcPr.remove(e)
    tcPr.append(shd)

def _borders(cell, color=AZUL, sz=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(sz))
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tcB.append(b)
    for e in tcPr.findall(qn('w:tcBorders')): tcPr.remove(e)
    tcPr.append(tcB)

def _pad(cell, top=80, bottom=80, left=120, right=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcM = OxmlElement('w:tcMar')
    for side, val in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), str(val))
        m.set(qn('w:type'), 'dxa')
        tcM.append(m)
    for e in tcPr.findall(qn('w:tcMar')): tcPr.remove(e)
    tcPr.append(tcM)

def _cw(cell, width):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    w = OxmlElement('w:tcW')
    w.set(qn('w:w'), str(width))
    w.set(qn('w:type'), 'dxa')
    for e in tcPr.findall(qn('w:tcW')): tcPr.remove(e)
    tcPr.append(w)

def _table_setup(table, col_widths):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(sum(col_widths)))
    tblW.set(qn('w:type'), 'dxa')
    for e in tblPr.findall(qn('w:tblW')): tblPr.remove(e)
    tblPr.append(tblW)
    lay = OxmlElement('w:tblLayout')
    lay.set(qn('w:type'), 'fixed')
    for e in tblPr.findall(qn('w:tblLayout')): tblPr.remove(e)
    tblPr.append(lay)
    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is None:
        tblGrid = OxmlElement('w:tblGrid')
        tbl.insert(1, tblGrid)
    else:
        for gc in tblGrid.findall(qn('w:gridCol')): tblGrid.remove(gc)
    for w in col_widths:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(w))
        tblGrid.append(gc)

# ─── Helpers de párrafo ───────────────────────────────────────────────────────

def _run(para, text, bold=False, color=NEGRO, size_pt=10):
    run = para.add_run(str(text or ''))
    run.font.name = 'Arial'
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    return run

def _clear(cell):
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)

def _para(cell, text='', bold=False, color=NEGRO, size_pt=10,
          align=None, sp_before=16, sp_after=16):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Twips(sp_before)
    p.paragraph_format.space_after  = Twips(sp_after)
    if align: p.alignment = align
    if text: _run(p, text, bold=bold, color=color, size_pt=size_pt)
    return p

def _lines(cell, text, size_pt=10, bold=False, color=NEGRO):
    _clear(cell)
    for line in str(text or '').split('\n'):
        _para(cell, line, bold=bold, color=color, size_pt=size_pt)

# ─── Tipos de celda ──────────────────────────────────────────────────────────

def _hdr(cell, text, width=None):
    _borders(cell); _pad(cell, 60, 60, 120, 120); _shading(cell, AZUL)
    if width: _cw(cell, width)
    _clear(cell)
    _para(cell, text, bold=True, color=BLANCO, size_pt=11,
          align=WD_ALIGN_PARAGRAPH.CENTER, sp_before=24, sp_after=24)

def _subhdr(cell, text, width=None):
    _borders(cell); _pad(cell); _shading(cell, AZUL_CLARO)
    if width: _cw(cell, width)
    _clear(cell)
    _para(cell, text, bold=True, color=AZUL,
          align=WD_ALIGN_PARAGRAPH.CENTER, sp_before=20, sp_after=20)

def _lbl(cell, text, width=None):
    _borders(cell); _pad(cell); _shading(cell, AZUL_CLARO)
    if width: _cw(cell, width)
    _clear(cell)
    _para(cell, text, bold=True, color=AZUL, sp_before=20, sp_after=20)

def _val(cell, text, width=None, size_pt=10, shade=None):
    _borders(cell); _pad(cell)
    if shade: _shading(cell, shade)
    if width: _cw(cell, width)
    _lines(cell, text, size_pt=size_pt)

def _val_s(cell, text, width=None):
    _val(cell, text, width=width, size_pt=9)

# ─── Utilidades de datos ──────────────────────────────────────────────────────

def _fmt_dur(minutos):
    m = int(minutos or 0)
    if m >= 60:
        h, r = divmod(m, 60)
        return f"{h} h {r} min" if r else f"{h} hrs"
    return f"{m} min"

def _tiempo(d, titulo):
    norm = lambda s: ' '.join(str(s or '').lower().split())
    for b in (d.get('tiempos') or []):
        for f in (b.get('filas') or []):
            if norm(f.get('titulo', '')) == norm(titulo):
                return f"{f.get('tiempo', '')} min"
    return ""

def _subtotal(d, seccion):
    for b in (d.get('tiempos') or []):
        if b.get('seccion') == seccion:
            return f"{sum(int(f.get('tiempo') or 0) for f in (b.get('filas') or []))} min"
    return "—"

def _inst(texto, fallback):
    if not texto: return fallback
    lines = [l.strip() for l in str(texto).split('\n') if l.strip()]
    return lines[0] if lines and len(lines[0]) <= 120 else fallback

# ─── Tablas ───────────────────────────────────────────────────────────────────

def _tabla_info_general(doc, d):
    C1 = round(TW * .20); C2 = round(TW * .30)
    C3 = round(TW * .18); C4 = TW - C1 - C2 - C3
    dat = d.get('datos') or {}

    t = doc.add_table(rows=0, cols=4)
    _table_setup(t, [C1, C2, C3, C4])

    r = t.add_row(); r.cells[0].merge(r.cells[3])
    _hdr(r.cells[0], "INFORMACIÓN GENERAL", TW)

    r = t.add_row(); r.cells[1].merge(r.cells[3])
    _lbl(r.cells[0], "Nombre del Curso / Sesión:", C1)
    _borders(r.cells[1]); _pad(r.cells[1]); _cw(r.cells[1], C2+C3+C4); _clear(r.cells[1])
    _para(r.cells[1], dat.get('nombreCurso',''), bold=True, size_pt=12, sp_before=20, sp_after=20)

    for lbl_txt, key in [("Nombre del Diseñador:", 'disenador'),
                          ("Nombre del Instructor / Facilitador:", 'instructor')]:
        r = t.add_row(); r.cells[1].merge(r.cells[3])
        _lbl(r.cells[0], lbl_txt, C1)
        _val(r.cells[1], dat.get(key, ''), C2+C3+C4)

    r = t.add_row()
    _lbl(r.cells[0], "Lugar de Instrucción:", C1); _val(r.cells[1], dat.get('lugar',''), C2)
    _lbl(r.cells[2], "Duración:", C3);            _val(r.cells[3], _fmt_dur(dat.get('duracion',0)), C4)

    r = t.add_row()
    _lbl(r.cells[0], "Fecha(s):", C1);            _val(r.cells[1], dat.get('fecha',''), C2)
    _lbl(r.cells[2], "Nº de participantes:", C3)
    part = dat.get('participantes')
    _val(r.cells[3], f"{part} participantes" if part else '', C4)

    r = t.add_row(); r.cells[1].merge(r.cells[3])
    _lbl(r.cells[0], "Perfil del participante:", C1)
    _val(r.cells[1], dat.get('perfil',''), C2+C3+C4)

    r = t.add_row(); r.cells[1].merge(r.cells[3])
    _lbl(r.cells[0], "Beneficios del curso:", C1)
    _val(r.cells[1], d.get('beneficios',''), C2+C3+C4, shade=AMARILLO)


def _tabla_objetivos(doc, d):
    C1 = round(TW * .65); C2 = TW - C1
    obj = d.get('objetivos') or {}; tem = d.get('temario') or {}

    t = doc.add_table(rows=0, cols=2)
    _table_setup(t, [C1, C2])

    DESCRIP_OG = "describe la demostración de un conocimiento, desempeño o producto, resultado del aprendizaje del participante, así como el dominio de aprendizaje cognitivo, psicomotriz, afectivo y relacional-social en los que impactará el Curso/sesión"

    r = t.add_row(); r.cells[0].merge(r.cells[1])
    _borders(r.cells[0]); _pad(r.cells[0]); _shading(r.cells[0], AZUL_CLARO); _clear(r.cells[0])
    p = r.cells[0].add_paragraph()
    p.paragraph_format.space_before = Twips(20); p.paragraph_format.space_after = Twips(20)
    _run(p, "Objetivo General ", bold=True, color=AZUL)
    _run(p, f"({DESCRIP_OG})", color=NEGRO, size_pt=9)

    r = t.add_row(); r.cells[0].merge(r.cells[1])
    _borders(r.cells[0]); _pad(r.cells[0]); _clear(r.cells[0])
    _para(r.cells[0], obj.get('general',''), sp_before=20, sp_after=20)

    r = t.add_row()
    _subhdr(r.cells[0], "Objetivos Particulares", C1)
    _subhdr(r.cells[1], "Temas:", C2)

    for num, tipo, key_obj, key_tem in [
        ("1", "(Cognitivo)",              'cognitiva',   'u1'),
        ("2", "(Psicomotriz)",            'psicomotriz', 'u2'),
        ("3", "(Afectivo / Relacional-social)", 'afectiva', 'u3'),
    ]:
        r = t.add_row()
        _borders(r.cells[0]); _pad(r.cells[0]); _cw(r.cells[0], C1); _clear(r.cells[0])
        p = r.cells[0].add_paragraph()
        p.paragraph_format.space_before = Twips(16); p.paragraph_format.space_after = Twips(8)
        _run(p, f"{num}. ", bold=True); _run(p, tipo, bold=True, color=AZUL)
        _para(r.cells[0], obj.get(key_obj,''), sp_before=8, sp_after=16)
        _val(r.cells[1], '\n'.join(tem.get(key_tem) or []), C2)


def _tabla_requerimientos(doc, d):
    CW = round(TW / 6); CW6 = TW - CW * 5
    mat = d.get('materiales') or {}
    vals = [
        mat.get('instalaciones') or mat.get('integracion')  or "Aula iluminada, ventilada, mesas y sillas suficientes.",
        mat.get('equipo')        or mat.get('expositiva')   or "Laptop, proyector, extensión eléctrica y presentación.",
        mat.get('materialesDidacticos') or mat.get('demostrativa') or "Manual del participante, hojas, bolígrafos y materiales de apoyo.",
        mat.get('humanos')       or mat.get('dialogo')      or "Instructor y participantes registrados.",
        mat.get('otros')         or mat.get('energizante')  or "Material extra requerido para las actividades.",
        mat.get('seguridad')     or "Botiquín, señalización de salida, medidas de higiene y protección civil.",
    ]
    labels = [
        "Instalaciones, mobiliario y distribución:",
        "Equipo de apoyo y distribución:",
        "Materiales didácticos de apoyo:",
        "Requerimientos humanos:",
        "Otros:",
        "Salud / Seguridad / Higiene / Protección civil:",
    ]
    widths = [CW, CW, CW, CW, CW, CW6]

    t = doc.add_table(rows=0, cols=6)
    _table_setup(t, widths)

    r = t.add_row(); r.cells[0].merge(r.cells[5])
    _hdr(r.cells[0], "REQUERIMIENTOS PARA EL DESARROLLO DEL CURSO", TW)

    r = t.add_row()
    for i, (lbl, w) in enumerate(zip(labels, widths)):
        _borders(r.cells[i]); _pad(r.cells[i], 60, 60, 80, 80)
        _shading(r.cells[i], AZUL_CLARO); _cw(r.cells[i], w); _clear(r.cells[i])
        _para(r.cells[i], lbl, bold=True, color=AZUL, size_pt=9, sp_before=16, sp_after=16)

    r = t.add_row()
    for i, (val, w) in enumerate(zip(vals, widths)):
        _borders(r.cells[i]); _pad(r.cells[i], 60, 60, 80, 80); _cw(r.cells[i], w)
        _lines(r.cells[i], val, size_pt=9)


def _tabla_evaluaciones(doc, d):
    ev = d.get('evaluaciones') or {}
    C1 = round(TW * .35); C2 = round(TW * .08)
    C3 = round(TW * .37); C4 = TW - C1 - C2 - C3
    pct_d = ev.get('pctDiagnostica') or ev.get('pctDiag') or 0
    pct_f = ev.get('pctFormativa')   or ev.get('pctForm') or 0
    pct_s = ev.get('pctSumativa')    or ev.get('pctSuma') or 0

    t = doc.add_table(rows=0, cols=4)
    _table_setup(t, [C1, C2, C3, C4])

    r = t.add_row()
    _subhdr(r.cells[0], "Aspecto a Evaluar / Finalidad", C1)
    _subhdr(r.cells[1], "%", C2)
    _subhdr(r.cells[2], "Instrumento de Evaluación", C3)
    _subhdr(r.cells[3], "Momento de Aplicación", C4)

    rows_data = [
        ("1. Evaluación Diagnóstica",
         "Identificar el nivel de conocimientos previos de los participantes como punto de partida del curso.",
         pct_d, _inst(ev.get('instDiagnostica'), 'Cuestionario diagnóstico'), "Al inicio\n(solo referencial)"),
        ("2. Evaluación Formativa",
         "Identificar la comprensión y avance logrado por los participantes durante el curso.",
         pct_f, _inst(ev.get('instFormativa'), 'Lista de cotejo / Guía de observación'), "Intermedia"),
        ("3. Evaluación Final (Sumativa)",
         "Acreditar los aprendizajes adquiridos por los participantes en el proceso de enseñanza-aprendizaje.",
         pct_s, _inst(ev.get('instSumativa'), 'Cuestionario final'), "Al final del curso"),
    ]
    for tipo, finalidad, pct, instrumento, momento in rows_data:
        r = t.add_row()
        _borders(r.cells[0]); _pad(r.cells[0]); _cw(r.cells[0], C1); _clear(r.cells[0])
        _para(r.cells[0], tipo, bold=True, sp_before=20, sp_after=8)
        _para(r.cells[0], f"Finalidad: {finalidad}", size_pt=9, sp_before=8, sp_after=20)
        _val(r.cells[1], f"{pct}%", C2); _val(r.cells[2], instrumento, C3); _val(r.cells[3], momento, C4)

    r = t.add_row(); r.cells[0].merge(r.cells[3])
    _borders(r.cells[0]); _pad(r.cells[0]); _shading(r.cells[0], GRIS_CLARO); _clear(r.cells[0])
    _para(r.cells[0], "d) Criterios de evaluación:", bold=True, sp_before=20, sp_after=8)
    for c in ["- Conocimientos Teóricos: Comprende los temas del curso, identifica conceptos clave y los relaciona con su práctica laboral.",
              "- Actitud y comportamiento: Participación activa, respeto a sus compañeros y al instructor.",
              "- Evaluaciones aplicadas: Comprensión de los temas explicados y puntaje aprobatorio en las evaluaciones efectuadas."]:
        _para(r.cells[0], c, size_pt=9, sp_before=6, sp_after=6)


def _act_paras(cell, text):
    lines = str(text or '').split('\n')
    for i, line in enumerate(lines):
        bold = i == 0 and bool(line.strip())
        _para(cell, line, bold=bold, color=AZUL if bold else NEGRO,
              size_pt=9, sp_before=16 if i == 0 else 6, sp_after=6)


def _tabla_seccion(doc, titulo, filas, primera_col="Etapa"):
    CE = round(TW * .16); CA = round(TW * .46)
    CD = round(TW * .07); CT = round(TW * .15); CM = TW - CE - CA - CD - CT

    t = doc.add_table(rows=0, cols=5)
    _table_setup(t, [CE, CA, CD, CT, CM])

    r = t.add_row(); r.cells[0].merge(r.cells[4])
    _hdr(r.cells[0], titulo, TW)

    r = t.add_row()
    _subhdr(r.cells[0], primera_col, CE); _subhdr(r.cells[1], "Actividades", CA)
    _subhdr(r.cells[2], "Duración", CD);  _subhdr(r.cells[3], "Técnicas Grupales / Instruccionales", CT)
    _subhdr(r.cells[4], "Material y Equipo de Apoyo", CM)

    skip = {'a)', 'b)', 'c)', 'd)', 'e)', 'f)'}
    for fila in filas:
        es_suma = 'suma' in str(fila.get('etapa', '')).lower()
        acts = [a for a in (fila.get('actividades') or []) if str(a or '').strip() not in skip and str(a or '').strip()]

        if es_suma:
            r = t.add_row(); r.cells[0].merge(r.cells[4])
            _borders(r.cells[0]); _pad(r.cells[0], 60, 60, 80, 80)
            _shading(r.cells[0], AZUL_CLARO); _clear(r.cells[0])
            for a in acts:
                _para(r.cells[0], a, bold=True, color=AZUL, size_pt=9, sp_before=20, sp_after=20)
        else:
            r = t.add_row()
            _borders(r.cells[0]); _pad(r.cells[0], 60, 60, 80, 80); _cw(r.cells[0], CE)
            _lines(r.cells[0], fila.get('etapa', ''), size_pt=9)
            _borders(r.cells[1]); _pad(r.cells[1], 60, 60, 80, 80); _cw(r.cells[1], CA); _clear(r.cells[1])
            for a in acts: _act_paras(r.cells[1], a)
            _val_s(r.cells[2], fila.get('duracion', ''), CD)
            _val_s(r.cells[3], fila.get('tecnica', ''), CT)
            _val_s(r.cells[4], fila.get('material', ''), CM)


def _tabla_desarrollo(doc, titulo, filas):
    _tabla_seccion(doc, titulo, filas, primera_col="Temas / Subtemas")


# ─── Constructores de filas ───────────────────────────────────────────────────

def _filas_previo(_d):
    return [{'etapa': "Comprobación de la existencia y funcionamiento de los recursos requeridos",
             'actividades': ["El instructor aplicará la Lista de verificación de requerimientos.",
                             "El instructor realizará pruebas de funcionamiento del equipo.",
                             "El instructor verificará la distribución del mobiliario y equipo.",
                             "El instructor verificará la suficiencia de los materiales conforme al número de participantes."],
             'duracion': "10 min previos al inicio", 'tecnica': "", 'material': "Lista de verificación de requerimientos"}]


def _filas_apertura(d):
    tec = d.get('tecnicas') or {}; enc = d.get('encuadre') or {}
    obj = d.get('objetivos') or {}; tem = d.get('temario') or {}; ev = d.get('evaluaciones') or {}
    rh  = tec.get('rhNombre', 'Técnica Rompe Hielo')
    rh_raw = tec.get('rhDetalle', '')
    rh_t = _tiempo(d, "Presentación del Instructor")
    partes = [p for p in rh_raw.split('\n\n') if not any(p.strip().startswith(x) for x in ('d)','e)','f)'))]
    partes = [p.replace('c) Mencionará el tiempo para realizarla.', f'c) Mencionará el tiempo para realizarla: {rh_t}') if p.strip().startswith('c)') and rh_t else p for p in partes]
    rh_det = '\n\n'.join(partes)
    reglas = (enc.get('reglasTexto') or []) + ([enc.get('otraRegla')] if enc.get('otraRegla') else [])
    pct_d = ev.get('pctDiagnostica') or ev.get('pctDiag') or 0
    pct_f = ev.get('pctFormativa')   or ev.get('pctForm') or 0
    pct_s = ev.get('pctSumativa')    or ev.get('pctSuma') or 0
    NL = '\n'
    return [
        {'etapa': "1. Presentación del instructor / facilitador y de los participantes",
         'actividades': ["El instructor se presentará ante el grupo.",
                         "El instructor pedirá que se anoten en la lista de asistencia.",
                         "El instructor propiciará la presentación de los participantes.",
                         f'El instructor aplicará la Técnica Rompe Hielo / Integración: "{rh}".\n{rh_det}'],
         'duracion': _tiempo(d, "Presentación del Instructor"), 'tecnica': f"Técnica grupal:\n{rh}", 'material': "Lista de asistencia\nMateriales de la técnica"},
        {'etapa': "2. Presentación del curso",
         'actividades': [
             f"El instructor presentará los objetivos del curso/sesión.\n\nOBJETIVO GENERAL:\n{obj.get('general','')}\n\nOBJETIVOS PARTICULARES:\n1. (Cognitivo) {obj.get('cognitiva','')}\n\n2. (Psicomotriz) {obj.get('psicomotriz','')}\n\n3. (Afectivo) {obj.get('afectiva','')}",
             f"El instructor presentará la descripción general del desarrollo del curso/sesión.\n{ev.get('descripcionGeneral','')}",
             f"El instructor mencionará el temario del curso/sesión.\n\nUnidad 1:\n{NL.join(tem.get('u1') or [])}\n\nUnidad 2:\n{NL.join(tem.get('u2') or [])}\n\nUnidad 3:\n{NL.join(tem.get('u3') or [])}",
             f"El instructor creará un ambiente participativo mediante preguntas.\n{enc.get('preguntas','')}",
             f"El instructor explicará los beneficios del curso/sesión.\n{d.get('beneficios','')}",
             f"El instructor especificará el tipo de evaluaciones a realizar.\n\na) Evaluación Diagnóstica ({_inst(ev.get('instDiagnostica'),'Cuestionario')}) — Al inicio. {pct_d}% solo referencial.\nb) Evaluación Formativa ({_inst(ev.get('instFormativa'),'Lista de cotejo')}) — Intermedia. {pct_f}%.\nc) Evaluación Sumativa ({_inst(ev.get('instSumativa'),'Cuestionario')}) — Al final. {pct_s}%.\nd) Criterios: Conocimientos Teóricos, Actitud y comportamiento, Evaluaciones aplicadas.",
         ], 'duracion': _tiempo(d, "Objetivos del curso/sesión"), 'tecnica': "", 'material': "Laptop, proyector, PowerPoint"},
        {'etapa': "3. Acuerdos y compromisos",
         'actividades': ["El instructor acordará con el grupo las expectativas del curso/sesión.",
                         f"El instructor acordará las reglas de operación.\n{NL.join(reglas)}",
                         "El instructor realizará el contrato de aprendizaje con los participantes."],
         'duracion': _tiempo(d, "Reglas de operación del curso"), 'tecnica': "", 'material': "Hojas blancas\nFormatos de contrato de aprendizaje"},
        {'etapa': "4. Evaluación diagnóstica",
         'actividades': ["El instructor realizará la evaluación diagnóstica.",
                         f"Instrumento: {_inst(ev.get('instDiagnostica'),'Cuestionario diagnóstico')}",
                         "Indicará alcance, propósito y finalidad.",
                         "Indicará las instrucciones y el tiempo para realizarla.",
                         "Aclarará las dudas que se presenten.",
                         "Mencionará que los errores son oportunidades para fortalecer el aprendizaje."],
         'duracion': _tiempo(d, "Evaluación diagnóstica"), 'tecnica': "", 'material': "Instrumentos diagnósticos\nBolígrafos"},
        {'etapa': "Suma de los tiempos",
         'actividades': [f"Subtotal Apertura / Encuadre: {_subtotal(d, 'Inicio / Encuadre del curso')}"],
         'duracion': _subtotal(d, "Inicio / Encuadre del curso"), 'tecnica': "", 'material': ""},
    ]


def _filas_desarrollo(d):
    tec = d.get('tecnicas') or {}; obj = d.get('objetivos') or {}
    tem = d.get('temario') or {}; ev  = d.get('evaluaciones') or {}
    exp = d.get('expositiva') or {}; dem = d.get('demostrativa') or {}; dia = d.get('dialogo') or {}
    en = tec.get('enNombre', 'Técnica Energizante'); en_det = tec.get('enDetalle', '')
    NL = '\n'
    return [
        {'etapa': f"Unidad 1 — Cognitivo\n\n{NL.join(tem.get('u1') or [])}",
         'actividades': ["El instructor aplicará la técnica expositiva:",
                         f"a) Presentará el objetivo del tema:\n{obj.get('cognitiva','')}",
                         f"b) Introducción general al contenido temático:\n{exp.get('introduccion','')}",
                         f"c) Recuperará la experiencia previa:\n{exp.get('experiencia','')}",
                         f"d) Desarrollará el contenido:\n{exp.get('desarrollo','')}",
                         f"e) Utilizará ejemplos:\n{exp.get('ejemplos','')}",
                         f"g) Realizará la síntesis:\n{exp.get('sintesis','')}",
                         f"h) Planteará preguntas dirigidas:\n{exp.get('preguntas','')}",
                         f"i) Promoverá comentarios sobre utilidad:\n{exp.get('utilidad','')}"],
         'duracion': _tiempo(d, "Técnica expositiva"), 'tecnica': "Expositiva", 'material': "Laptop, proyector, PowerPoint"},
        {'etapa': f"Unidad 2 — Psicomotriz\n\n{NL.join(tem.get('u2') or [])}",
         'actividades': ["El instructor aplicará la técnica demostrativa:",
                         f"a) Presentará objetivo de la actividad:\n{obj.get('psicomotriz','')}",
                         f"b) Recuperará la experiencia previa:\n{dem.get('experiencia','')}",
                         f"c) Presentará la actividad a desarrollar:\n{dem.get('actividad','')}",
                         "d) Ejemplificará la actividad.", "e) Resolverá dudas.", "f) Permitirá la práctica.", "g) Retroalimentará.",
                         f"h) Usará ejemplos:\n{dem.get('ejemplos','')}",
                         f"i) Preguntará por los conocimientos adquiridos:\n{dem.get('preguntas','')}",
                         "j) Recordará las reglas de operación.", "k) Mencionará los logros alcanzados."],
         'duracion': _tiempo(d, "Técnica demostrativa"), 'tecnica': "Demostrativa", 'material': "Laptop, proyector, PowerPoint"},
        {'etapa': "Evaluación Formativa",
         'actividades': ["El instructor realizará la evaluación formativa.",
                         f"Instrumento: {_inst(ev.get('instFormativa'),'Lista de cotejo / Guía de observación')}",
                         "Indicará alcance, propósito y finalidad.", "Indicará instrucciones y tiempo.", "Aclarará dudas."],
         'duracion': _tiempo(d, "Evaluación formativa"), 'tecnica': "", 'material': "Formatos de evaluación formativa\nBolígrafos"},
        {'etapa': "Descanso", 'actividades': ["Descanso"],
         'duracion': _tiempo(d, "Descanso"), 'tecnica': "", 'material': "Servicio de café"},
        {'etapa': f"Técnica Energizante\n\n{en}",
         'actividades': [f'El instructor aplicará la Técnica Energizante: "{en}".',
                         f"a) Explicará objetivo:\n{tec.get('enObjetivo','')}",
                         f"b) Dará instrucciones:\n{en_det}",
                         "c) Mencionará el tiempo.", "d) Participará con el grupo.", "e) Controlará el tiempo."],
         'duracion': _tiempo(d, "Técnica Energizante"), 'tecnica': f"Técnica grupal:\n{en}", 'material': "Materiales de la técnica"},
        {'etapa': f"Unidad 3 — Afectivo / Relacional-social\n\n{NL.join(tem.get('u3') or [])}",
         'actividades': ["El instructor aplicará la técnica diálogo-discusión:",
                         f"a) Mencionará el objetivo:\n{obj.get('afectiva','')}",
                         f"b) Presentará la actividad:\n{dia.get('actividad','')}",
                         f"c) Mencionará el tema a discutir:\n{dia.get('tema','')}",
                         f"d) Indicará las instrucciones:\n{dia.get('instrucciones','')}",
                         f"e) Indicará el tiempo:\n{dia.get('tiempo','')}", "f) Dividirá al grupo en subgrupos.",
                         f"g) Establecerá reglas:\n{dia.get('reglas','')}",
                         f"h) Abrirá la discusión:\n{dia.get('introduccion','')}",
                         "i) Propiciará la discusión.", "j) Moderará la discusión.",
                         f"k) Utilizará ejemplos:\n{dia.get('ejemplos','')}",
                         f"l) Desarrollará conclusiones:\n{dia.get('conclusion','')}"],
         'duracion': _tiempo(d, "Técnica diálogo-discusión"), 'tecnica': "Diálogo-discusión", 'material': "Laptop, proyector, PowerPoint"},
        {'etapa': "Evaluación Final",
         'actividades': ["El instructor realizará la evaluación final.",
                         f"Instrumento: {_inst(ev.get('instSumativa'),'Cuestionario final')}",
                         "Indicará alcances e instrucciones.", "Indicará el tiempo.", "Aclarará dudas."],
         'duracion': _tiempo(d, "Evaluación final"), 'tecnica': "", 'material': "Formatos de evaluación final\nBolígrafos"},
        {'etapa': "Suma de los tiempos",
         'actividades': [f"Subtotal Desarrollo: {_subtotal(d, 'Desarrollo del curso')}"],
         'duracion': _subtotal(d, "Desarrollo del curso"), 'tecnica': "", 'material': ""},
    ]


def _filas_cierre(d):
    obj = d.get('objetivos') or {}; c = d.get('cierre') or {}
    total = sum(int(f.get('tiempo') or 0) for b in (d.get('tiempos') or []) for f in (b.get('filas') or []))
    return [
        {'etapa': "1. Conclusiones",
         'actividades': [a for a in ["El instructor realizará la conclusión de los contenidos temáticos desarrollados.",
                                     "Mencionará los logros alcanzados.",
                                     "Preguntará la opinión sobre la aplicación de los temas.",
                                     c.get('texto','')] if a],
         'duracion': _tiempo(d, "Conclusión"), 'tecnica': "Técnica grupal de cierre", 'material': "Laptop, proyector, PowerPoint"},
        {'etapa': "2. Resumen general del curso",
         'actividades': [a for a in ["El instructor mencionará el resumen general del curso.", c.get('resumen','')] if a],
         'duracion': _tiempo(d, "Resumen general del curso"), 'tecnica': "", 'material': ""},
        {'etapa': "3. Logro de expectativas",
         'actividades': ["El instructor retomará las expectativas escritas al inicio y analizará si se cumplieron."],
         'duracion': _tiempo(d, "Logro de expectativas del curso"), 'tecnica': "", 'material': "Pintarrón"},
        {'etapa': "4. Logro de objetivos",
         'actividades': [f"El instructor preguntará acerca del logro de los objetivos.\n\nOBJETIVO GENERAL:\n{obj.get('general','')}\n\n¿Se cumplió el objetivo general del curso?"],
         'duracion': _tiempo(d, "Logro de los objetivos"), 'tecnica': "", 'material': "Laptop, proyector, PowerPoint"},
        {'etapa': "5. Sugerencias de continuidad",
         'actividades': [a for a in ["El instructor sugerirá acciones de continuidad en el aprendizaje.", c.get('sugerencias','')] if a],
         'duracion': _tiempo(d, "Sugerencias de continuidad del aprendizaje"), 'tecnica': "", 'material': ""},
        {'etapa': "6. Referencia(s) bibliográfica(s)",
         'actividades': [a for a in ["El instructor indicará las referencias bibliográficas.", c.get('referencias','')] if a],
         'duracion': _tiempo(d, "Referencia(s) bibliográfica"), 'tecnica': "", 'material': ""},
        {'etapa': "7. Compromisos de aplicación",
         'actividades': [a for a in ["El instructor conducirá al grupo a formular compromisos de aplicación.", c.get('compromisos','')] if a],
         'duracion': _tiempo(d, "Compromisos de aplicación del aprendizaje"), 'tecnica': "", 'material': ""},
        {'etapa': "8. Evaluación de satisfacción",
         'actividades': ["El instructor aplicará la evaluación de satisfacción.",
                         "Indicará alcances e instrucciones de la evaluación.", "Indicará el tiempo.", "Aclarará dudas."],
         'duracion': _tiempo(d, "Evaluación de satisfacción"), 'tecnica': "", 'material': "Instrumentos de evaluación de satisfacción\nBolígrafos"},
        {'etapa': "9. Cierre",
         'actividades': ["El instructor empleará una técnica de cierre.", "El instructor dará las gracias."],
         'duracion': _tiempo(d, "Cierre"), 'tecnica': "", 'material': ""},
        {'etapa': "Suma de los tiempos",
         'actividades': [f"Subtotal Cierre: {_subtotal(d, 'Cierre del curso')}", f"TOTAL GENERAL: {total} min"],
         'duracion': _subtotal(d, "Cierre del curso"), 'tecnica': "", 'material': ""},
    ]


# ─── Función principal ─────────────────────────────────────────────────────────

def generar_planeacion_docx_v2(payload: dict, branding: dict | None = None) -> bytes:
    br = branding or {}
    empresa_footer = br.get("empresa") or "SmartBuilder EC  •  Centro ECM"
    logo_url = br.get("logo_url")

    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Twips(12240); sec.page_height = Twips(15840)
    sec.left_margin = sec.right_margin = Twips(720)
    sec.top_margin  = sec.bottom_margin = Twips(720)
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(10)

    # D#7: Logo del admin si existe
    if logo_url:
        try:
            import urllib.request, tempfile, os
            with urllib.request.urlopen(logo_url, timeout=5) as resp:
                logo_bytes = resp.read()
            ext = logo_url.split(".")[-1].lower().split("?")[0]
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}")
            tmp.write(logo_bytes); tmp.close()
            lp = doc.add_paragraph()
            lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            lp.paragraph_format.space_before = Twips(0)
            lp.paragraph_format.space_after  = Twips(80)
            run = lp.add_run()
            run.add_picture(tmp.name, height=Pt(36))
            os.unlink(tmp.name)
        except Exception as e_logo:
            print(f"[branding] No se pudo insertar logo: {e_logo}")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Twips(0)
    p.paragraph_format.space_after  = Twips(240)
    r = p.add_run("DOCUMENTO DE PLANEACIÓN DEL CURSO / CARTA DESCRIPTIVA")
    r.font.name = 'Arial'; r.font.size = Pt(14); r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(AZUL)

    _tabla_info_general(doc, payload)
    doc.add_paragraph().paragraph_format.space_before = Twips(200)
    _tabla_objetivos(doc, payload)
    doc.add_paragraph().paragraph_format.space_before = Twips(200)
    _tabla_requerimientos(doc, payload)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Twips(80)
    p.paragraph_format.space_after  = Twips(80)
    r = p.add_run("Formas, momentos y criterios de evaluación: La evaluación se llevará a cabo durante la Apertura, el Desarrollo y el Cierre del Curso/sesión.")
    r.font.name = 'Arial'; r.font.size = Pt(11); r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(AZUL)
    _tabla_evaluaciones(doc, payload)

    doc.add_page_break()

    _tabla_seccion(doc, "PREVIO AL INICIO DEL CURSO — Comprobación de Recursos", _filas_previo(payload))
    doc.add_paragraph().paragraph_format.space_before = Twips(200)
    _tabla_seccion(doc, "INICIO DEL CURSO — APERTURA O ENCUADRE", _filas_apertura(payload))
    doc.add_page_break()
    _tabla_desarrollo(doc, "DESARROLLO", _filas_desarrollo(payload))
    doc.add_page_break()
    _tabla_seccion(doc, "CIERRE", _filas_cierre(payload))

    dat = payload.get('datos') or {}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Twips(320)
    r = p.add_run(f"{empresa_footer}  •  {dat.get('nombreCurso','')}  •  {dat.get('fecha','')}")
    r.font.name = 'Arial'; r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string("999999")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
